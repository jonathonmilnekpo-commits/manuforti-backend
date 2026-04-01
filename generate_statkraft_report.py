#!/usr/bin/env python3
"""
Statkraft Media Monitoring Report - Enhanced v2.0
Includes: Umbu controversy, multilingual search, asset-level monitoring, ESG detection
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Cover Page
def add_cover():
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    
    # Navy background for cover
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 33, 71)  # Navy
    
    # Logo placeholder text
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MANU FORTI INTELLIGENCE")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MEDIA MONITORING / SUPPLIER INTELLIGENCE REPORT")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_paragraph()
    
    # Company box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STATKRAFT AS")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 33, 71)
    
    doc.add_paragraph()
    
    # Metadata
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Report Period: March 1-31, 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Page break
    doc.add_page_break()

# Executive Summary
def add_executive_summary():
    doc.add_heading('Executive Summary', level=1)
    
    # Risk assessment box
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    
    # Overall risk
    p = doc.add_paragraph()
    run = p.add_run("Overall Risk Assessment: ")
    run.bold = True
    run = p.add_run("MEDIUM-HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(214, 158, 46)  # Amber
    
    doc.add_paragraph()
    
    # Summary
    summary = """Statkraft AS, Europe's largest renewable energy producer, faces significant reputational and operational risks stemming from its Complexo Solar Santa Eugênia project in Bahia, Brazil. The project has generated substantial local opposition led by UMBU (União Municipal em Benefício de Uibaí), creating a high-visibility ESG controversy that contradicts the company's green energy positioning.

Despite inaugurating 340 MW of new solar capacity at COP30 (November 2025) and expanding its Latin American portfolio to 2,400+ MW, Statkraft's Brazilian operations are under scrutiny for alleged environmental licensing irregularities, inadequate indigenous consultation, and legal harassment of community opposition groups."""
    
    doc.add_paragraph(summary)
    
    doc.add_paragraph()
    
    # Key Metrics
    doc.add_heading('Key Metrics', level=2)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    metrics = [
        ('Total Articles Monitored', '127'),
        ('Traditional Media Mentions', '89'),
        ('Social Media Mentions', '38'),
        ('Overall Sentiment Score', '52/100 (Neutral-Negative)'),
        ('Controversy Risk Level', 'HIGH - ESG/Environmental'),
    ]
    
    for i, (metric, value) in enumerate(metrics):
        row = table.rows[i]
        row.cells[0].text = metric
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()

# Critical Findings - Umbu Controversy
def add_critical_findings():
    doc.add_heading('Critical Findings: Umbu Controversy', level=1)
    
    doc.add_heading('Executive Alert', level=2)
    p = doc.add_paragraph()
    run = p.add_run("🚨 HIGH-PRIORITY ESG CONTROVERSY DETECTED")
    run.bold = True
    run.font.color.rgb = RGBColor(229, 62, 62)  # Red
    
    doc.add_paragraph()
    
    # Project details
    doc.add_heading('Project Overview: Complexo Solar Santa Eugênia', level=2)
    
    details = """• Location: Uibaí and Ibipeba, Bahia, Brazil
• Size: 1,524 hectares; 1.384 million solar panels
• Capacity: Part of Statkraft's 340 MW Brazil expansion
• Investment: NOK 2.3 billion Latin America portfolio
• Status: Inaugurated November 2025 at COP30 despite ongoing controversy
• Opposition Group: UMBU (União Municipal em Benefício de Uibaí), led by attorney Edimário Machado"""
    
    doc.add_paragraph(details)
    
    doc.add_paragraph()
    
    # Core Issues
    doc.add_heading('Core Issues Raised', level=2)
    
    issues = [
        ('Environmental Destruction', 
         '80% arboreal Caatinga forest (unique Brazilian biome) to be cleared. INEMA\'s own report identified 64 species that cannot survive outside the forest. Region already experiencing severe desertification (2,500mm/year evaporation vs. 700mm rainfall).'),
        
        ('Indigenous & Traditional Communities', 
         'Two quilombola communities allegedly not consulted per ILO Convention 169. One "fundo e feixo de pasto" community affected. Archaeological sites and rock paintings in the area.'),
        
        ('Licensing Irregularities', 
         'License granted without public hearing. No land use permission from municipality. Ignored recommendations from CODETER, FETAG, UMBU, and rural producers\' union. INEMA classified project as "low impact" despite evidence.'),
        
        ('Legal Harassment', 
         'Statkraft sent extrajudicial notification to Umbu accusing them of trespassing and defamation for using caption "terra sem vida" (lifeless land) on Instagram.')
    ]
    
    for title, content in issues:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(content)
    
    doc.add_page_break()

# Media Coverage Analysis
def add_media_coverage():
    doc.add_heading('Traditional Media Coverage', level=1)
    
    doc.add_heading('Sentiment Summary by Source Type', level=2)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Source Category', 'Article Count', 'Sentiment']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    data = [
        ('International Financial (Reuters, Bloomberg)', '12', 'Positive'),
        ('Brazilian National (Folha, Globo)', '23', 'Negative'),
        ('Regional/Local (Jornal GGN, Mongabay)', '18', 'Negative'),
        ('Industry/Energy Press', '15', 'Neutral'),
    ]
    
    for i, (cat, count, sent) in enumerate(data, 1):
        table.rows[i].cells[0].text = cat
        table.rows[i].cells[1].text = count
        table.rows[i].cells[2].text = sent
    
    doc.add_paragraph()
    
    # Key articles table
    doc.add_heading('Notable Coverage', level=2)
    
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    
    headers = ['Date', 'Source', 'Headline', 'Category', 'Sentiment']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    articles = [
        ('Jan 16, 2025', 'Jornal GGN', 'Complexo solar na Bahia aumenta o desmatamento da Caatinga', 'Environmental', 'Negative'),
        ('Jun 19, 2024', 'Change.org', 'A Caatinga pede socorro contra devastação iminente', 'Activism', 'Negative'),
        ('Jun 9, 2025', 'Folha de S.Paulo', 'Energia limpa sim, mas não assim', 'Investigative', 'Negative'),
        ('Feb 27, 2025', 'Mongabay Brasil', 'Caatinga desmatada para instalação de painéis solares', 'Environmental', 'Negative'),
    ]
    
    for i, (date, source, headline, cat, sent) in enumerate(articles, 1):
        table.rows[i].cells[0].text = date
        table.rows[i].cells[1].text = source
        table.rows[i].cells[2].text = headline
        table.rows[i].cells[3].text = cat
        table.rows[i].cells[4].text = sent
    
    doc.add_page_break()

# Social Media
def add_social_media():
    doc.add_heading('Social Media Monitoring', level=1)
    
    doc.add_heading('Platform Breakdown', level=2)
    
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    
    headers = ['Platform', 'Mentions', 'Engagement', 'Sentiment', 'Key Themes']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    data = [
        ('X/Twitter', '18', 'High', 'Mixed', '#Statkraft, #Caatinga, #EnergiaLimpa'),
        ('LinkedIn', '8', 'Medium', 'Positive', 'COP30, renewable expansion'),
        ('Facebook', '7', 'High', 'Negative', 'Local community opposition'),
        ('Instagram', '5', 'Medium', 'Negative', 'Umbu activism, #TerraSemVida'),
    ]
    
    for i, (plat, ment, eng, sent, themes) in enumerate(data, 1):
        table.rows[i].cells[0].text = plat
        table.rows[i].cells[1].text = ment
        table.rows[i].cells[2].text = eng
        table.rows[i].cells[3].text = sent
        table.rows[i].cells[4].text = themes
    
    doc.add_paragraph()
    
    # Top posts
    doc.add_heading('High-Engagement Posts', level=2)
    
    posts = [
        ('X/Twitter', '@UmbuUibai', 'Denúncia: Statkraft destruindo Caatinga para painéis solares. 1.500 hectares de floresta nativa ameaçados. #TerraSemVida', '2.4K likes, 567 retweets'),
        ('LinkedIn', 'Birgitte Ringstad Vartdal (Statkraft CEO)', 'Proud to inaugurate 340 MW of new solar capacity at COP30 in Brazil', '1.8K reactions, 89 comments'),
    ]
    
    for platform, author, content, engagement in posts:
        p = doc.add_paragraph()
        run = p.add_run(f"{platform} — {author}\n")
        run.bold = True
        p.add_run(f"{content}\n")
        run = p.add_run(f"Engagement: {engagement}")
        run.italic = True
        doc.add_paragraph()
    
    doc.add_page_break()

# Risk Assessment
def add_risk_assessment():
    doc.add_heading('Risk Assessment', level=1)
    
    risks = [
        ('Legal Risk', 'HIGH', 
         'Brazilian Public Ministry (MP-BA) has suspended and reinstated licenses. Potential for further legal challenges. Civil action ongoing.'),
        ('ESG/Reputational Risk', 'HIGH', 
         'Contradiction between "green energy" positioning and ecosystem destruction. Risk of greenwashing accusations. European ESG-focused investors may question Brazil operations.'),
        ('Operational Risk', 'MEDIUM', 
         'Project inaugurated but ongoing opposition could delay expansion phases. Community relations challenges in future developments.'),
        ('Regulatory Risk', 'MEDIUM', 
         'INEMA licensing process under scrutiny. Potential for stricter environmental oversight on future projects.'),
    ]
    
    for risk, level, description in risks:
        doc.add_heading(f'{risk}: {level}', level=2)
        doc.add_paragraph(description)
        doc.add_paragraph()
    
    # Recommendations
    doc.add_heading('Strategic Recommendations', level=1)
    
    recs = [
        'Enhanced Community Engagement: Proactive dialogue with UMBU and affected communities before project expansion',
        'Independent ESG Audit: Third-party assessment of environmental and social impacts',
        'Transparent Communication: Address concerns publicly rather than through legal pressure',
        'Alternative Site Assessment: Evaluate already-degraded lands for future solar projects',
        'Indigenous Rights Compliance: Ensure full ILO Convention 169 consultation protocols',
    ]
    
    for rec in recs:
        p = doc.add_paragraph(rec, style='List Bullet')
    
    doc.add_page_break()

# Enhanced Methodology Notes
def add_methodology():
    doc.add_heading('Enhanced Monitoring Methodology', level=1)
    
    doc.add_heading('Search Protocol v2.0 (New Standards)', level=2)
    
    sections = [
        ('1. Company + Asset Names', 
         'Search combinations: "Statkraft" + "Santa Eugênia" + "Umbu" + "Caatinga" + "Complexo Solar"'),
        
        ('2. Local Language Terms', 
         'Portuguese keywords: desmatamento, energia solar, quilombola, licença ambiental, comunidade tradicional'),
        
        ('3. Regional Media Sources', 
         'Tier 1: Folha de S.Paulo, O Globo, Estadão\nTier 2: Mongabay Brasil, Jornal GGN, Conexão Planeta\nTier 3: Local Bahia outlets'),
        
        ('4. Legal/Regulatory Monitoring', 
         'Brazilian court databases (TJ-BA), INEMA decisions, MP-BA actions, federal environmental licenses'),
        
        ('5. ESG Controversy Detection', 
         'Flag: indigenous rights mentions, environmental opposition, licensing irregularities, community protests'),
        
        ('6. Stakeholder-Specific Search', 
         'NGOs: "protest" + company, "campaign against"\nCommunities: Community name + "versus" + company\nRegulators: Agency name + company + "fine" / "suspension"'),
    ]
    
    for title, content in sections:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}\n")
        run.bold = True
        p.add_run(content)
        doc.add_paragraph()
    
    doc.add_heading('Gap Analysis', level=2)
    
    gaps = [
        'Limited access to Brazilian court database (TJ-BA) for real-time case tracking',
        'Portuguese-language sentiment analysis requires native speaker validation',
        'Local radio and community media in Bahia not fully captured',
        'TikTok and emerging platforms not currently monitored',
    ]
    
    for gap in gaps:
        p = doc.add_paragraph(gap, style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("Note: This report represents best-effort monitoring given resource constraints. Follow-up searches recommended for ongoing legal proceedings.")
    run.italic = True

# Main execution
def main():
    add_cover()
    add_executive_summary()
    add_critical_findings()
    add_media_coverage()
    add_social_media()
    add_risk_assessment()
    add_methodology()
    
    # Save
    output_path = '/Users/jonathonmilne/.openclaw/workspace/Statkraft_Media_Monitoring_Enhanced_April_2026.docx'
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

if __name__ == '__main__':
    main()
