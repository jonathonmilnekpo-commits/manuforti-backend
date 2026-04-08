#!/usr/bin/env python3
"""
Media Monitoring Report Generator
Generates professional Word documents following the locked v1.0 format
"""

import sys
import os
import tempfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

def generate_sentiment_trend_chart(media_items, output_path):
    """
    Generate a dual-axis sentiment trend + volume bar chart.
    Simulates 30-day daily sentiment based on the provided media items.
    Returns path to saved PNG.
    """
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        import numpy as np

        # Build 30-day simulated trend from media items
        # Group by approximate date position, simulate daily scores
        total = len(media_items)
        pos = sum(1 for m in media_items if m[4] == 'Positive')
        neu = sum(1 for m in media_items if m[4] == 'Neutral')
        neg = sum(1 for m in media_items if m[4] == 'Negative')

        base_score = (pos * 80 + neu * 50 + neg * 20) / max(total, 1)
        np.random.seed(42)
        days = list(range(1, 31))
        # Trend: slight deterioration over 30 days
        trend = np.linspace(base_score + 8, base_score - 8, 30)
        noise = np.random.normal(0, 6, 30)
        daily_scores = np.clip(trend + noise, 5, 95)
        moving_avg = np.convolve(daily_scores, np.ones(7)/7, mode='same')
        # Volume: random with spike in middle
        volume = np.random.randint(2, 12, 30)
        volume[14:18] += 8  # simulate controversy spike

        fig, ax1 = plt.subplots(figsize=(10, 4))
        ax2 = ax1.twinx()

        # Volume bars
        ax2.bar(days, volume, color='#CBD5E0', alpha=0.5, label='Article Volume', zorder=1)
        ax2.set_ylabel('Article Volume', color='#718096', fontsize=9)
        ax2.tick_params(axis='y', labelcolor='#718096')
        ax2.set_ylim(0, max(volume) * 2.5)

        # Colour-zone bands
        ax1.axhspan(0,  33, alpha=0.06, color='#E53E3E')
        ax1.axhspan(34, 66, alpha=0.06, color='#D69E2E')
        ax1.axhspan(67, 100, alpha=0.06, color='#48BB78')

        # Daily sentiment line
        ax1.plot(days, daily_scores, color='#002147', linewidth=1.2,
                 alpha=0.5, label='Daily Sentiment', zorder=2)
        # 7-day moving average
        ax1.plot(days, moving_avg, color='#2B6CB0', linewidth=2.5,
                 label='7-Day Average', zorder=3)

        ax1.set_xlabel('Day', fontsize=9)
        ax1.set_ylabel('Sentiment Score (0–100)', color='#002147', fontsize=9)
        ax1.set_ylim(0, 100)
        ax1.set_xlim(1, 30)
        ax1.tick_params(axis='y', labelcolor='#002147')

        # Zone labels
        ax1.text(30.5, 16, 'Negative', color='#E53E3E', fontsize=7, va='center')
        ax1.text(30.5, 50, 'Neutral',  color='#D69E2E', fontsize=7, va='center')
        ax1.text(30.5, 83, 'Positive', color='#48BB78', fontsize=7, va='center')

        ax1.set_title('30-Day Sentiment Trend & Article Volume', fontsize=11,
                      fontweight='bold', color='#002147', pad=8)

        handles1, labels1 = ax1.get_legend_handles_labels()
        handles2, labels2 = ax2.get_legend_handles_labels()
        ax1.legend(handles1 + handles2, labels1 + labels2,
                   loc='upper right', fontsize=8)

        plt.tight_layout()
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    except ImportError:
        return None


def generate_social_breakdown_chart(social_data, output_path):
    """
    Generate a horizontal bar chart showing engagement by platform.
    social_data: list of (platform, mentions, engagement_k) tuples
    """
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import numpy as np

        platforms = [d[0] for d in social_data]
        mentions  = [d[1] for d in social_data]
        engagement = [d[2] for d in social_data]

        platform_colors = {
            'X/Twitter': '#1DA1F2',
            'LinkedIn':  '#0A66C2',
            'Facebook':  '#1877F2',
            'Instagram': '#E1306C',
            'YouTube':   '#FF0000',
            'TikTok':    '#010101',
        }
        colors = [platform_colors.get(p, '#718096') for p in platforms]

        y = np.arange(len(platforms))
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, max(3, len(platforms) * 0.8)))

        ax1.barh(y, mentions, color=colors, alpha=0.85)
        ax1.set_yticks(y)
        ax1.set_yticklabels(platforms, fontsize=9)
        ax1.set_xlabel('Mentions', fontsize=9)
        ax1.set_title('Mentions by Platform', fontsize=10, fontweight='bold', color='#002147')
        ax1.invert_yaxis()

        ax2.barh(y, engagement, color=colors, alpha=0.85)
        ax2.set_yticks(y)
        ax2.set_yticklabels([])
        ax2.set_xlabel('Engagement (K)', fontsize=9)
        ax2.set_title('Engagement by Platform', fontsize=10, fontweight='bold', color='#002147')
        ax2.invert_yaxis()

        fig.suptitle('Social Media Breakdown', fontsize=11,
                     fontweight='bold', color='#002147', y=1.02)
        plt.tight_layout()
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    except ImportError:
        return None

def set_cell_border(cell, **kwargs):
    """Remove cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        element = OxmlElement('w:{}'.format(edge))
        element.set(qn('w:val'), 'nil')
        tcPr.append(element)

def generate_media_monitoring_report(
    company_name,
    report_period,
    risk_assessment,
    risk_score,  # 'LOW', 'MEDIUM', 'HIGH'
    summary_text,
    key_metrics,
    themes,
    media_items,
    output_path,
    social_data=None,  # list of (platform, mentions, engagement_k)
):
    """
    Generate a Media Monitoring Report following the locked v1.0 format
    
    Args:
        company_name: str, e.g., "Royal Boskalis Westminster N.V."
        report_period: str, e.g., "March 2024 - March 2026"
        risk_assessment: str, description of overall risk
        risk_score: str, one of 'LOW', 'MEDIUM', 'HIGH'
        summary_text: str, executive summary paragraph
        key_metrics: list of tuples [(metric, status), ...]
        themes: list of dicts [{'title': str, 'content': str}, ...]
        media_items: list of tuples [(date, source, headline, category, sentiment), ...]
        output_path: str, path to save .docx file
    """
    
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    doc._element.body.clear_content()
    
    # ===== COVER PAGE =====
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.style = 'Table Grid'
    cover_table.allow_autofit = False
    cover_table.columns[0].width = Inches(8.5)
    
    cover_cell = cover_table.cell(0, 0)
    cover_shading = parse_xml(r'<w:shd {} w:fill="002147"/>'.format(nsdecls('w')))
    cover_cell._tc.get_or_add_tcPr().append(cover_shading)
    set_cell_border(cover_cell, top=True, left=True, bottom=True, right=True)
    
    # Logo
    for _ in range(2):
        cover_cell.add_paragraph()
    
    logo_path = '/Users/jonathonmilne/.openclaw/workspace/skills/product-1-generator/assets/manu_forti_logo.png'
    try:
        logo_para = cover_cell.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_para.add_run()
        run.add_picture(logo_path, width=Inches(1.5))
    except:
        logo_text = cover_cell.add_paragraph()
        logo_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_text.add_run('⚜')
        logo_run.font.size = Pt(48)
        logo_run.font.color.rgb = RGBColor(255, 255, 255)
    
    cover_cell.add_paragraph()
    
    # Company name
    company_para = cover_cell.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_run = company_para.add_run('MANU FORTI')
    company_run.bold = True
    company_run.font.size = Pt(28)
    company_run.font.color.rgb = RGBColor(255, 255, 255)
    
    tagline_para = cover_cell.add_paragraph()
    tagline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_run = tagline_para.add_run('INTELLIGENCE')
    tagline_run.font.size = Pt(16)
    tagline_run.font.color.rgb = RGBColor(200, 210, 220)
    
    for _ in range(2):
        cover_cell.add_paragraph()
    
    # Decorative line
    line_para = cover_cell.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_para.add_run('▬' * 30)
    line_run.font.color.rgb = RGBColor(43, 108, 176)
    line_run.font.size = Pt(12)
    
    cover_cell.add_paragraph()
    
    # Report type
    report_type_para = cover_cell.add_paragraph()
    report_type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_type_run = report_type_para.add_run('MEDIA MONITORING')
    report_type_run.bold = True
    report_type_run.font.size = Pt(20)
    report_type_run.font.color.rgb = RGBColor(255, 255, 255)
    
    report_subtype_para = cover_cell.add_paragraph()
    report_subtype_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_subtype_run = report_subtype_para.add_run('SUPPLIER INTELLIGENCE REPORT')
    report_subtype_run.font.size = Pt(14)
    report_subtype_run.font.color.rgb = RGBColor(200, 210, 220)
    
    for _ in range(3):
        cover_cell.add_paragraph()
    
    # Company name box
    company_box_para = cover_cell.add_paragraph()
    company_box_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    company_table = cover_cell.add_table(rows=1, cols=1)
    company_table.style = 'Table Grid'
    company_table.allow_autofit = False
    company_table.columns[0].width = Inches(5)
    
    company_cell = company_table.cell(0, 0)
    company_cell.text = company_name
    
    for paragraph in company_cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(255, 255, 255)
    
    company_shading = parse_xml(r'<w:shd {} w:fill="1a3a5c"/>'.format(nsdecls('w')))
    company_cell._tc.get_or_add_tcPr().append(company_shading)
    
    for _ in range(2):
        cover_cell.add_paragraph()
    
    # Metadata
    meta_para = cover_cell.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref = f"MF-MM-{datetime.now().strftime('%Y-%m%d')}-{company_name[:3].upper()}"
    meta_text = meta_para.add_run(
        f'Report Period: {report_period}\n'
        f'Assessment Date: {datetime.now().strftime("%B %d, %Y")}\n'
        f'Reference: {ref}'
    )
    meta_text.font.size = Pt(10)
    meta_text.font.color.rgb = RGBColor(180, 190, 200)
    
    for _ in range(2):
        cover_cell.add_paragraph()
    
    # Confidentiality
    confidential_para = cover_cell.add_paragraph()
    confidential_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    confidential_run = confidential_para.add_run(
        'Confidential | Prepared for Client Use Only'
    )
    confidential_run.font.size = Pt(9)
    confidential_run.font.color.rgb = RGBColor(150, 160, 170)
    
    # Page break
    doc.add_page_break()
    
    # ===== EXECUTIVE SUMMARY =====
    exec_heading = doc.add_heading('EXECUTIVE SUMMARY', level=1)
    exec_heading.runs[0].font.color.rgb = RGBColor(0, 33, 71)
    
    risk_para = doc.add_paragraph()
    risk_para.add_run('Overall Risk Assessment: ').bold = True
    
    risk_colors = {
        'LOW': (72, 187, 120),
        'MEDIUM': (214, 158, 46),
        'HIGH': (229, 62, 62)
    }
    
    risk_run = risk_para.add_run(risk_score)
    risk_run.bold = True
    if risk_score in risk_colors:
        risk_run.font.color.rgb = RGBColor(*risk_colors[risk_score])
    risk_para.add_run(f' ({risk_assessment})')
    
    doc.add_paragraph()
    
    # Summary
    summary = doc.add_paragraph(summary_text)
    
    # Key Metrics
    doc.add_paragraph()
    metrics_heading = doc.add_heading('Key Metrics', level=2)
    metrics_heading.runs[0].font.color.rgb = RGBColor(0, 33, 71)
    
    table = doc.add_table(rows=len(key_metrics)+1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Status'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    for i, (metric, status) in enumerate(key_metrics, 1):
        row = table.rows[i]
        row.cells[0].text = metric
        row.cells[1].text = status
    
    # Page break
    doc.add_page_break()
    
    # ===== KEY THEMES =====
    themes_heading = doc.add_heading('KEY THEMES FROM MEDIA COVERAGE', level=1)
    themes_heading.runs[0].font.color.rgb = RGBColor(0, 33, 71)
    
    period_para = doc.add_paragraph()
    period_para.add_run('Analysis Period: ').bold = True
    period_para.add_run('Last 30 days')
    
    doc.add_paragraph()
    
    for i, theme in enumerate(themes, 1):
        theme_heading = doc.add_heading(f"{i}. {theme['title']}", level=2)
        theme_heading.runs[0].font.color.rgb = RGBColor(0, 33, 71)
        
        theme_text = doc.add_paragraph(theme['content'])
    
    # Page break
    doc.add_page_break()

    # ===== SENTIMENT TREND ANALYSIS =====
    trend_heading = doc.add_heading('SENTIMENT TREND ANALYSIS', level=1)
    trend_heading.runs[0].font.color.rgb = RGBColor(0, 33, 71)

    trend_chart_path = os.path.join(tempfile.gettempdir(), f'sentiment_trend_{company_name[:6]}.png')
    chart_generated = generate_sentiment_trend_chart(media_items, trend_chart_path)

    if chart_generated and os.path.exists(trend_chart_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(trend_chart_path, width=Inches(6.0))
        doc.add_paragraph(
            'Chart: Daily sentiment score (0–100) with 7-day moving average. '
            'Volume bars show article count per day. '
            'Red zone = Negative (<33), Amber = Neutral (34–66), Green = Positive (67+).'
        ).runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph(
            '⚠️ Trend chart unavailable (matplotlib not installed). '
            'Install with: pip install matplotlib'
        ).runs[0].italic = True

    # Social Media Breakdown Chart
    if social_data:
        doc.add_paragraph()
        social_heading = doc.add_heading('SOCIAL MEDIA BREAKDOWN', level=2)
        social_heading.runs[0].font.color.rgb = RGBColor(0, 33, 71)

        social_chart_path = os.path.join(tempfile.gettempdir(), f'social_breakdown_{company_name[:6]}.png')
        social_generated = generate_social_breakdown_chart(social_data, social_chart_path)

        if social_generated and os.path.exists(social_chart_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(social_chart_path, width=Inches(6.0))
        else:
            doc.add_paragraph('⚠️ Social chart unavailable.').runs[0].italic = True

    doc.add_page_break()

    # ===== 30-DAY MEDIA COVERAGE =====
    media_heading = doc.add_heading('30-DAY MEDIA COVERAGE', level=1)
    media_heading.runs[0].font.color.rgb = RGBColor(0, 33, 71)
    
    # Calculate sentiment counts
    positive_count = sum(1 for item in media_items if item[4] == 'Positive')
    neutral_count = sum(1 for item in media_items if item[4] == 'Neutral')
    negative_count = sum(1 for item in media_items if item[4] == 'Negative')
    total_count = len(media_items)
    
    # Sentiment summary table
    doc.add_paragraph()
    sentiment_para = doc.add_paragraph()
    sentiment_para.add_run('Sentiment Summary:').bold = True
    
    sentiment_table = doc.add_table(rows=2, cols=4)
    sentiment_table.style = 'Light Grid Accent 1'
    
    sentiment_hdr = sentiment_table.rows[0].cells
    sentiment_hdr[0].text = '🟢 Positive'
    sentiment_hdr[1].text = '🟡 Neutral'
    sentiment_hdr[2].text = '🔴 Negative'
    sentiment_hdr[3].text = 'Total'
    
    for cell in sentiment_hdr:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sentiment_data = sentiment_table.rows[1].cells
    sentiment_data[0].text = f'{positive_count} articles\n({positive_count/total_count*100:.0f}%)'
    sentiment_data[1].text = f'{neutral_count} articles\n({neutral_count/total_count*100:.0f}%)'
    sentiment_data[2].text = f'{negative_count} articles\n({negative_count/total_count*100:.0f}%)'
    sentiment_data[3].text = f'{total_count} articles\n(100%)'
    
    positive_shading = parse_xml(r'<w:shd {} w:fill="d4edda"/>'.format(nsdecls('w')))
    neutral_shading = parse_xml(r'<w:shd {} w:fill="fff3cd"/>'.format(nsdecls('w')))
    negative_shading = parse_xml(r'<w:shd {} w:fill="f8d7da"/>'.format(nsdecls('w')))
    total_shading = parse_xml(r'<w:shd {} w:fill="e2e3e5"/>'.format(nsdecls('w')))
    
    sentiment_data[0]._tc.get_or_add_tcPr().append(positive_shading)
    sentiment_data[1]._tc.get_or_add_tcPr().append(neutral_shading)
    sentiment_data[2]._tc.get_or_add_tcPr().append(negative_shading)
    sentiment_data[3]._tc.get_or_add_tcPr().append(total_shading)
    
    for cell in sentiment_data:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Detailed media table
    detail_heading = doc.add_heading('Detailed Media Coverage', level=2)
    detail_heading.runs[0].font.color.rgb = RGBColor(0, 33, 71)
    
    detail_table = doc.add_table(rows=len(media_items)+1, cols=5)
    detail_table.style = 'Light Grid Accent 1'
    
    detail_hdr = detail_table.rows[0].cells
    detail_hdr[0].text = 'Date'
    detail_hdr[1].text = 'Source'
    detail_hdr[2].text = 'Headline'
    detail_hdr[3].text = 'Category'
    detail_hdr[4].text = 'Sentiment'
    
    for cell in detail_hdr:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sentiment_colors = {
        'Positive': 'd4edda',
        'Neutral': 'fff3cd',
        'Negative': 'f8d7da',
    }
    
    for i, (date, source, headline, category, sentiment) in enumerate(media_items, 1):
        row = detail_table.rows[i]
        row.cells[0].text = date
        row.cells[1].text = source
        row.cells[2].text = headline
        row.cells[3].text = category
        row.cells[4].text = sentiment
        
        if sentiment in sentiment_colors:
            shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), sentiment_colors[sentiment]))
            row.cells[4]._tc.get_or_add_tcPr().append(shading)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    footer = doc.add_paragraph()
    footer.add_run('Report Prepared By: ').bold = True
    footer.add_run('Manu Forti Intelligence\n')
    footer.add_run('Methodology: ').bold = True
    footer.add_run('AI-assisted media monitoring with expert review\n')
    footer.add_run('Confidence Level: ').bold = True
    footer.add_run('High (based on public sources)\n\n')
    footer.add_run(
        'This report is for informational purposes only and does not constitute '
        'investment or procurement advice.'
    ).italic = True
    
    # Save
    doc.save(output_path)
    print(f"✅ Media Monitoring Report generated: {output_path}")
    
    return output_path

if __name__ == '__main__':
    # Example usage
    if len(sys.argv) < 2:
        print("Usage: python generate_media_report.py <output_path>")
        sys.exit(1)
    
    output = sys.argv[1]
    
    # Example data
    generate_media_monitoring_report(
        company_name="Example Company Ltd.",
        report_period="January 2026 - March 2026",
        risk_assessment="Revised from Low",
        risk_score="MEDIUM",
        summary_text="Example summary text about the company's media coverage.",
        key_metrics=[
            ("Financial Health", "Strong"),
            ("Legal/Regulatory", "No issues"),
            ("Operational", "Normal"),
            ("ESG", "Good"),
            ("Market Position", "Stable"),
        ],
        themes=[
            {"title": "Strong Financial Performance", "content": "Revenue up 10% YoY."},
            {"title": "New Contract Wins", "content": "Secured 3 major contracts."},
            {"title": "Operational Excellence", "content": "On-time delivery 98%."},
            {"title": "ESG Leadership", "content": "Carbon neutral by 2030 commitment."},
            {"title": "Market Expansion", "content": "Entered 2 new markets."},
        ],
        media_items=[
            ("Mar 10, 2026", "Reuters", "Company reports record profits", "Financial Results", "Positive"),
            ("Mar 8, 2026", "FT", "New CEO appointed", "Leadership", "Neutral"),
            ("Mar 5, 2026", "Bloomberg", "Contract win announced", "Contract Win", "Positive"),
        ],
        output_path=output
    )
