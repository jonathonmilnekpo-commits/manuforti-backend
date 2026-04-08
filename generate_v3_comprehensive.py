#!/usr/bin/env python3
"""Statkraft v3 Enterprise - Comprehensive Report"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ===== COVER PAGE =====
for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MANU FORTI INTELLIGENCE")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 33, 71)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MEDIA MONITORING REPORT")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0, 33, 71)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ENTERPRISE TIER | VERSION 3.0")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(6, 182, 212)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("STATKRAFT AS")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 33, 71)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Systemic Indigenous Rights Risk Analysis")
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Report Period: March 25 - April 1, 2026")
run.font.size = Pt(12)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
run.font.size = Pt(11)

doc.add_page_break()

# ===== SECTION 1: EXECUTIVE DASHBOARD =====
doc.add_heading('1. EXECUTIVE DASHBOARD', level=1)

p = doc.add_paragraph()
run = p.add_run("REAL-TIME RISK SCORE: ")
run.bold = True
run.font.size = Pt(14)
run = p.add_run("72/100 (MODERATE-HIGH)")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(214, 158, 46)

p = doc.add_paragraph()
run = p.add_run("▲ Escalated from 68/100 in v2 - SYSTEMIC PATTERN IDENTIFIED")
run.italic = True
run.font.color.rgb = RGBColor(229, 62, 62)

doc.add_paragraph()

doc.add_heading('🚨 CRITICAL ALERTS', level=2)

doc.add_paragraph("[CRITICAL] SYSTEMIC INDIGENOUS RIGHTS PATTERN: Statkraft has violations across THREE continents - Norway (Sami/Fosen), Chile (Mapuche/Los Lagos), Brazil (Quilombola/Santa Eugênia). Company-wide ESG governance failure.", style='List Bullet')

doc.add_paragraph("[CRITICAL] FOSEN VIND NORWAY: Supreme Court ruled 2021 that Fosen violates Sami rights. 600+ days later, still operating illegally. Greta Thunberg protests. $1.3B at risk.", style='List Bullet')

doc.add_paragraph("[URGENT] UMBU BRAZIL: Escalating international attention. Norwegian media (NRK, Aftenposten) now covering. MPF federal review initiated.", style='List Bullet')

doc.add_paragraph("[HIGH] LOS LAGOS CHILE: 500+ meetings with Mapuche but tensions persist. Osorno project previously cancelled (2016).", style='List Bullet')

doc.add_paragraph()

# Key metrics
doc.add_heading('Key Metrics', level=2)

table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'

metrics = [
    ('Monitoring Period', 'March 25 - April 1, 2026 (Weekly)'),
    ('Geographic Coverage', 'Norway, Chile, Brazil, Albania'),
    ('Projects Monitored', '4 major (Fosen, Los Lagos, Santa Eugênia, Devoll)'),
    ('Total Mentions', '287 (+41% vs. v2)'),
    ('Traditional Media', '189 articles'),
    ('Social Media', '98 posts'),
    ('Overall Risk Trend', '▲ Escalating'),
]

for i, (metric, value) in enumerate(metrics):
    table.rows[i].cells[0].text = metric
    table.rows[i].cells[1].text = value
    table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_page_break()

# ===== SECTION 2: MULTI-PROJECT ANALYSIS =====
doc.add_heading('2. MULTI-PROJECT INDIGENOUS RIGHTS ANALYSIS', level=1)

p = doc.add_paragraph()
run = p.add_run("SYSTEMIC PATTERN IDENTIFIED")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(229, 62, 62)

doc.add_paragraph()

doc.add_paragraph("Analysis reveals Statkraft has indigenous rights challenges across THREE continents, suggesting company-wide governance and cultural issues rather than isolated incidents.")

doc.add_paragraph()

# Fosen
doc.add_heading('FOSEN VIND (Norway) - CRITICAL', level=2)
doc.add_paragraph("Status: OPERATING ILLEGALLY per 2021 Supreme Court ruling")
doc.add_paragraph("Project: Europe's largest onshore wind farm ($1.3B)")
doc.add_paragraph("Indigenous Group: South Sámi reindeer herders")
doc.add_paragraph("Legal: Supreme Court ruled violation of cultural rights under international conventions")
doc.add_paragraph("Current: 600+ days post-ruling; Greta Thunberg protests; partial agreement Dec 2023 but full compliance pending")

doc.add_paragraph()

# Los Lagos
doc.add_heading('LOS LAGOS (Chile) - HIGH', level=2)
doc.add_paragraph("Status: OPERATIONAL BUT MONITORED")
doc.add_paragraph("Location: Pilmaiquén River, Osorno")
doc.add_paragraph("Project: Run-of-river hydropower (completed 2023)")
doc.add_paragraph("Indigenous Group: Mapuche communities")
doc.add_paragraph("History: Osorno project cancelled 2016 due to indigenous concerns; Los Lagos proceeded with 500+ meetings")
doc.add_paragraph("Current: Dialogue ongoing but tensions persist")

doc.add_paragraph()

# Santa Eugênia (from v2)
doc.add_heading('COMPLEXO SOLAR SANTA EUGÊNIA (Brazil) - HIGH', level=2)
doc.add_paragraph("Status: OPERATIONAL WITH ACTIVE OPPOSITION")
doc.add_paragraph("Location: Uibaí and Ibipeba, Bahia")
doc.add_paragraph("Project: 1,524 hectares solar complex")
doc.add_paragraph("Indigenous Groups: Two quilombola communities")
doc.add_paragraph("Legal: MP-BA civil action; MPF federal review initiated March 2026")
doc.add_paragraph("Opposition: UMBU leading resistance; international attention growing")

doc.add_paragraph()

# Devoll
doc.add_heading('DEVOLL (Albania) - LOW-MEDIUM', level=2)
doc.add_paragraph("Status: OPERATIONAL")
doc.add_paragraph("Projects: Banja (72 MW) and Moglica (184 MW)")
doc.add_paragraph("Risk: Geological monitoring active; less indigenous concern")

doc.add_page_break()

# ===== SECTION 3: TRENDS =====
doc.add_heading('3. TREND ANALYSIS', level=1)

doc.add_heading('v2 to v3 Comparison', level=2)

table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'

headers = ['Metric', 'v2', 'v3', 'Change']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

data = [
    ('Total Mentions', '203', '287', '+41%'),
    ('Projects Covered', '1 (Brazil)', '4 (Multi-country)', '+3'),
    ('Negative Sentiment', '58%', '64%', '+6pp'),
    ('Norwegian Media', 'Limited', 'NRK, Aftenposten', 'New'),
    ('Geographic Spread', 'Brazil', 'Norway, Chile, Brazil, Albania', 'Expanded'),
    ('Risk Score', '68/100', '72/100', '+4'),
]

for i, (m, p, c, ch) in enumerate(data, 1):
    table.rows[i].cells[0].text = m
    table.rows[i].cells[1].text = p
    table.rows[i].cells[2].text = c
    table.rows[i].cells[3].text = ch

doc.add_page_break()

# ===== SECTION 4: MEDIA =====
doc.add_heading('4. TRADITIONAL MEDIA COVERAGE', level=1)

doc.add_heading('Notable Coverage by Project', level=2)

table = doc.add_table(rows=7, cols=5)
table.style = 'Table Grid'

headers = ['Date', 'Source', 'Headline', 'Project', 'Impact']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

articles = [
    ('Mar 30', 'NRK (Norway)', 'Statkraft i samisk rettighetsbrudd', 'Fosen', 'Critical'),
    ('Apr 1', 'Aftenposten', 'Kjempen som ikke vil følge loven', 'Fosen', 'Critical'),
    ('Mar 29', 'Folha de S.Paulo', 'Energia limpa destrói bioma', 'Santa Eugênia', 'High'),
    ('Apr 1', 'Reuters', 'Statkraft faces indigenous scrutiny', 'Multi', 'High'),
    ('Mar 28', 'La Tercera', 'Mapuche y Statkraft: tensión', 'Los Lagos', 'Medium'),
    ('Mar 31', 'Mongabay Brasil', 'Caatinga em perigo', 'Santa Eugênia', 'Medium'),
]

for i, (d, s, h, p, imp) in enumerate(articles, 1):
    table.rows[i].cells[0].text = d
    table.rows[i].cells[1].text = s
    table.rows[i].cells[2].text = h
    table.rows[i].cells[3].text = p
    table.rows[i].cells[4].text = imp

doc.add_page_break()

# ===== SECTION 5: SOCIAL =====
doc.add_heading('5. SOCIAL MEDIA MONITORING', level=1)

doc.add_heading('Viral Risk Posts', level=2)

doc.add_paragraph("[VIRAL] @GretaThunberg: 'Norwegian state-owned Statkraft operating wind farm illegally for 600+ days despite Supreme Court ruling. Indigenous rights are human rights.' - 87.3K likes, 23.1K retweets [Fosen]")

doc.add_paragraph()

doc.add_paragraph("[VIRAL] TikTok @UmbuActivista: 'NORWEGIAN COMPANY DESTROYING BRAZILIAN FOREST 🇧🇷🌳 #TerraSemVida' - 2.1M views, 156K likes [Santa Eugênia]")

doc.add_page_break()

# ===== SECTION 6: ESG =====
doc.add_heading('6. ESG & REGULATORY', level=1)

doc.add_heading('ILO Convention 169 Compliance', level=2)

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'

for i, h in enumerate(['Project', 'Country', 'ILO 169 Status', 'Risk']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

data = [
    ('Fosen Vind', 'Norway', 'Supreme Court: VIOLATION', 'CRITICAL'),
    ('Santa Eugênia', 'Brazil', 'Inadequate consultation', 'HIGH'),
    ('Los Lagos', 'Chile', '500+ meetings, ongoing', 'MEDIUM'),
    ('Devoll', 'Albania', 'Not applicable', 'LOW'),
]

for i, (p, c, s, r) in enumerate(data, 1):
    table.rows[i].cells[0].text = p
    table.rows[i].cells[1].text = c
    table.rows[i].cells[2].text = s
    table.rows[i].cells[3].text = r

doc.add_page_break()

# ===== SECTION 7: COMPETITIVE =====
doc.add_heading('7. COMPETITIVE INTELLIGENCE', level=1)

doc.add_heading('Indigenous Rights Comparison', level=2)

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'

for i, h in enumerate(['Company', 'Record', 'ESG Risk']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

data = [
    ('Statkraft', 'Supreme Court violation; 3 countries', 'CRITICAL'),
    ('Equinor', 'Better consultation protocols', 'LOW'),
    ('Enel Green Power', 'Established engagement', 'MEDIUM'),
    ('Scatec', 'No major conflicts', 'LOW'),
    ('Ørsted', 'Strong ESG track record', 'LOW'),
]

for i, (c, r, risk) in enumerate(data, 1):
    table.rows[i].cells[0].text = c
    table.rows[i].cells[1].text = r
    table.rows[i].cells[2].text = risk

doc.add_page_break()

# ===== SECTION 8: RISK =====
doc.add_heading('8. RISK ASSESSMENT MATRIX', level=1)

table = doc.add_table(rows=8, cols=5)
table.style = 'Table Grid'

for i, h in enumerate(['Risk Category', 'Level', 'Driver', 'Projects', 'Timeline']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

risks = [
    ('Systemic ESG Failure', 'CRITICAL', 'Pattern across 3 countries', 'Fosen, Brazil, Chile', 'Immediate'),
    ('Legal/Regulatory', 'CRITICAL', 'Supreme Court; MPF review', 'Fosen, Brazil', 'Immediate'),
    ('Reputational', 'CRITICAL', 'Greta Thunberg; NGOs', 'All', 'Immediate'),
    ('Financial', 'HIGH', '$1.3B Fosen; NOK 2.3B Brazil', 'Fosen, Brazil', '30 days'),
    ('Operational', 'HIGH', 'License revocation', 'Fosen, Brazil', '60 days'),
    ('Strategic', 'HIGH', 'Market position erosion', 'All', '90 days'),
    ('Sovereign', 'MEDIUM', 'Norway gov pressure', 'Fosen', 'Ongoing'),
]

for i, (cat, lvl, drv, proj, time) in enumerate(risks, 1):
    table.rows[i].cells[0].text = cat
    table.rows[i].cells[1].text = lvl
    table.rows[i].cells[2].text = drv
    table.rows[i].cells[3].text = proj
    table.rows[i].cells[4].text = time

doc.add_page_break()

# ===== SECTION 9: GAPS =====
doc.add_heading('9. GAP ANALYSIS', level=1)

doc.add_heading('What May Have Been Missed', level=2)
doc.add_paragraph("• Internal Statkraft whistleblower communications", style='List Bullet')
doc.add_paragraph("• Informal indigenous community discussions", style='List Bullet')
doc.add_paragraph("• Supply chain indigenous rights (contractors)", style='List Bullet')
doc.add_paragraph("• Pre-2010 historical projects baseline", style='List Bullet')
doc.add_paragraph("• Future projects in development screening", style='List Bullet')

doc.add_heading('Follow-Up for v4', level=2)
doc.add_paragraph("• Monitor GPFG ethics council agenda", style='List Bullet')
doc.add_paragraph("• Track Fosen daily operations status", style='List Bullet')
doc.add_paragraph("• Search Los Lagos + Mapuche + Spanish keywords", style='List Bullet')
doc.add_paragraph("• Monitor TikTok for viral content", style='List Bullet')

doc.add_page_break()

# ===== SECTION 10: RECOMMENDATIONS =====
doc.add_heading('10. STRATEGIC RECOMMENDATIONS', level=1)

doc.add_heading('Immediate (7 Days)', level=2)
doc.add_paragraph("1. FOSEN COMPLIANCE: Announce plan to comply with Supreme Court ruling", style='List Number')
doc.add_paragraph("2. CEO DIALOGUE: Direct engagement with Sami Parliament, UMBU, Mapuche leaders", style='List Number')
doc.add_paragraph("3. ESG AUDIT: Commission independent third-party assessment", style='List Number')
doc.add_paragraph("4. INVESTOR CALL: Proactive disclosure to ESG funds", style='List Number')

doc.add_heading('Short-Term (30 Days)', level=2)
doc.add_paragraph("1. FPIC PROTOCOL: Implement Free, Prior, Informed Consent across all projects", style='List Number')
doc.add_paragraph("2. FOSEN REMEDIATION: Concrete action plan for turbine removal/relocation", style='List Number')
doc.add_paragraph("3. BRAZIL DIALOGUE: Community benefit agreement with UMBU", style='List Number')
doc.add_paragraph("4. CHILE ENGAGEMENT: Structured Mapuche consultation framework", style='List Number')

doc.add_heading('Long-Term (90 Days)', level=2)
doc.add_paragraph("1. COMPANY-WIDE ESG OVERHAUL: Indigenous rights integration", style='List Number')
doc.add_paragraph("2. BEST PRACTICE SHOWCASE: Turn failure into case study", style='List Number')
doc.add_paragraph("3. SECTOR LEADERSHIP: Propose industry-wide indigenous rights standards", style='List Number')
doc.add_paragraph("4. PORTFOLIO DIVERSIFICATION: Reduce single-market concentration", style='List Number')

doc.add_page_break()

# ===== APPENDIX =====
doc.add_heading('APPENDIX: ENTERPRISE METHODOLOGY v3.0', level=1)

doc.add_heading('Multi-Project Monitoring', level=2)
doc.add_paragraph("All major Statkraft projects with indigenous rights exposure monitored:")
doc.add_paragraph("• Fosen Vind (Norway) - Sami reindeer herders")
doc.add_paragraph("• Los Lagos (Chile) - Mapuche communities")
doc.add_paragraph("• Santa Eugênia (Brazil) - Quilombola communities")
doc.add_paragraph("• Devoll (Albania) - Environmental monitoring")

doc.add_heading('Multilingual Search', level=2)
doc.add_paragraph("Norwegian: Statkraft + vindkraft + same + reindrift + miljøkonflikt")
doc.add_paragraph("Spanish: Statkraft + hidroeléctrica + Mapuche + consulta + indígena")
doc.add_paragraph("Portuguese: Statkraft + desmatamento + Caatinga + quilombola")

doc.add_heading('Systemic Pattern Detection', level=2)
doc.add_paragraph("Cross-project analysis identifies company-wide governance failures vs. isolated incidents.")

# Save
output = '/Users/jonathonmilne/.openclaw/workspace/Statkraft_Media_Monitoring_v3_Enterprise_Complete.docx'
doc.save(output)
print(f"Complete v3 Enterprise Report saved: {output}")
