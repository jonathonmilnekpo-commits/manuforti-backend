#!/usr/bin/env python3
"""Statkraft v4 Complete - All 10 Sections"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# COVER
for _ in range(2): doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MANU FORTI INTELLIGENCE\nMEDIA MONITORING REPORT")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 33, 71)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ENTERPRISE TIER | VERSION 4.0")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(6, 182, 212)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("STATKRAFT AS")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 33, 71)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Report Period: April 1-8, 2026\nBuilds on v3 (March 25-April 1)")
run.font.size = Pt(11)

doc.add_page_break()

# VERSION CONTINUITY
doc.add_heading('VERSION CONTINUITY', level=1)
doc.add_paragraph("v4 BUILDS ON v3: All previous findings preserved and extended")
doc.add_heading('Carried Forward from v3', level=2)
doc.add_paragraph("• Fosen Vind (Norway) - Supreme Court violation", style='List Bullet')
doc.add_paragraph("• Los Lagos (Chile) - Mapuche dialogue", style='List Bullet')
doc.add_paragraph("• Santa Eugênia (Brazil) - UMBU opposition", style='List Bullet')
doc.add_paragraph("• Devoll (Albania) - Environmental monitoring", style='List Bullet')
doc.add_paragraph("• Systemic indigenous rights pattern", style='List Bullet')
doc.add_heading('NEW in v4', level=2)
doc.add_paragraph("• Norway PM compliance announcement", style='List Bullet')
doc.add_paragraph("• UN Special Rapporteur inquiry", style='List Bullet')
doc.add_paragraph("• Brazil MPF preliminary findings", style='List Bullet')
doc.add_paragraph("• Moody's ESG rating under review", style='List Bullet')
doc.add_page_break()

# SECTION 1
doc.add_heading('1. EXECUTIVE DASHBOARD', level=1)
p = doc.add_paragraph()
run = p.add_run("RISK SCORE: 71/100 (MODERATE-HIGH) ▼ from 72/100 in v3")
run.bold = True
run.font.size = Pt(14)
doc.add_paragraph()
doc.add_heading('New Developments', level=2)
doc.add_paragraph("[IMPROVING] Norway PM announced accelerated Fosen compliance", style='List Bullet')
doc.add_paragraph("[NEW] UN Special Rapporteur inquiry opened", style='List Bullet')
doc.add_paragraph("[ESCALATING] Brazil MPF finds INEMA 'systematically inadequate'", style='List Bullet')
doc.add_paragraph("[NEW] Moody's ESG rating under review", style='List Bullet')
doc.add_page_break()

# SECTION 2
doc.add_heading('2. CRITICAL FINDINGS', level=1)
doc.add_heading('Fosen Vind (Norway) - IMPROVING', level=2)
doc.add_paragraph("v3: Operating illegally 600+ days")
doc.add_paragraph("v4: PM Støre announced Q3 2026 compliance timeline; Sami Parliament cautiously optimistic")
doc.add_heading('Santa Eugênia (Brazil) - ESCALATING', level=2)
doc.add_paragraph("v3: MPF review initiated")
doc.add_paragraph("v4: MPF preliminary findings suggest federal oversight for all Statkraft Brazil projects")
doc.add_heading('NEW: UN Inquiry', level=2)
doc.add_paragraph("UN Special Rapporteur on Indigenous Rights opened formal inquiry (April 3, 2026)")
doc.add_paragraph("Scope: Pattern across Norway, Chile, Brazil | Report: Q3 2026")
doc.add_page_break()

# SECTION 3
doc.add_heading('3. TREND ANALYSIS', level=1)
doc.add_heading('4-Week Trend', level=2)
table = doc.add_table(rows=6, cols=5)
table.style = 'Table Grid'
for i, h in enumerate(['Metric', 'v1', 'v2', 'v3', 'v4']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
data = [
    ('Risk Score', '52/100', '68/100', '72/100', '71/100'),
    ('Projects', '1', '1', '4', '4'),
    ('Mentions', '127', '203', '287', '342'),
    ('Geography', 'Brazil', 'Brazil', '4 countries', '4+UN'),
    ('Legal Cases', '1', '2', '3', '4'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val
doc.add_page_break()

# SECTION 4
doc.add_heading('4. TRADITIONAL MEDIA', level=1)
doc.add_heading('New Coverage Since v3', level=2)
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
for i, h in enumerate(['Date', 'Source', 'Headline', 'Impact']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
articles = [
    ('Apr 5', 'NRK', 'Støre: Fosen løsning innen Q3', 'High'),
    ('Apr 4', 'Folha', 'MPF: Licenciamento systematically inadequate', 'Critical'),
    ('Apr 3', 'Reuters', 'UN opens Statkraft indigenous inquiry', 'Critical'),
    ('Apr 2', 'Aftenposten', 'Moody\'s setter Statkraft ESG under review', 'High'),
    ('Apr 2', 'La Tercera', 'Los Lagos: monitoreo ambiental cumple', 'Medium'),
]
for i, (d, s, h, imp) in enumerate(articles, 1):
    table.rows[i].cells[0].text = d
    table.rows[i].cells[1].text = s
    table.rows[i].cells[2].text = h
    table.rows[i].cells[3].text = imp
doc.add_page_break()

# SECTION 5
doc.add_heading('5. SOCIAL MEDIA', level=1)
doc.add_heading('Platform Activity', level=2)
table = doc.add_table(rows=5, cols=5)
table.style = 'Table Grid'
for i, h in enumerate(['Platform', 'Fosen', 'Brazil', 'Chile', 'Engagement']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
data = [
    ('X/Twitter', '24', '18', '10', '45.2K'),
    ('LinkedIn', '12', '8', '5', '11.3K'),
    ('Facebook', '8', '15', '7', '14.1K'),
    ('TikTok', '3', '8', '2', '3.8M'),
]
for i, (p, f, b, c, e) in enumerate(data, 1):
    table.rows[i].cells[0].text = p
    table.rows[i].cells[1].text = f
    table.rows[i].cells[2].text = b
    table.rows[i].cells[3].text = c
    table.rows[i].cells[4].text = e
doc.add_page_break()

# SECTION 6
doc.add_heading('6. ESG & REGULATORY', level=1)
doc.add_heading('New Legal/Regulatory Developments', level=2)
doc.add_paragraph("NORWAY: PM announced compliance acceleration (improving)")
doc.add_paragraph("BRAZIL: MPF preliminary findings critical (escalating)")
doc.add_paragraph("UN: Special Rapporteur inquiry opened (new)")
doc.add_paragraph("ESG: Moody's rating under review (new)")
doc.add_heading('ILO 169 Status Update', level=2)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
for i, h in enumerate(['Project', 'v3 Status', 'v4 Status', 'Trend']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
data = [
    ('Fosen', 'Supreme Court violation', 'PM compliance plan announced', '▲'),
    ('Santa Eugênia', 'MPF review initiated', 'MPF findings critical', '▼'),
    ('Los Lagos', '500+ meetings ongoing', 'Environmental compliance confirmed', '▲'),
    ('Devoll', 'Operational', 'Operational', '→'),
]
for i, (p, v3, v4, t) in enumerate(data, 1):
    table.rows[i].cells[0].text = p
    table.rows[i].cells[1].text = v3
    table.rows[i].cells[2].text = v4
    table.rows[i].cells[3].text = t
doc.add_page_break()

# SECTION 7
doc.add_heading('7. COMPETITIVE INTELLIGENCE', level=1)
doc.add_heading('ESG Rating Comparison', level=2)
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
for i, h in enumerate(['Company', 'ESG Rating', 'Indigenous Track Record', 'Status']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
data = [
    ('Statkraft', 'Under Review (Moody\'s)', 'Supreme Court violation; 3 countries', 'Deteriorating'),
    ('Equinor', 'A- (stable)', 'Better consultation protocols', 'Stable'),
    ('Enel', 'BBB+ (stable)', 'Established engagement', 'Stable'),
    ('Scatec', 'A- (stable)', 'No major conflicts', 'Stable'),
    ('Ørsted', 'A (stable)', 'Strong ESG track record', 'Stable'),
]
for i, (c, r, tr, s) in enumerate(data, 1):
    table.rows[i].cells[0].text = c
    table.rows[i].cells[1].text = r
    table.rows[i].cells[2].text = tr
    table.rows[i].cells[3].text = s
doc.add_page_break()

# SECTION 8
doc.add_heading('8. RISK ASSESSMENT MATRIX', level=1)
table = doc.add_table(rows=8, cols=5)
table.style = 'Table Grid'
for i, h in enumerate(['Risk', 'v3 Level', 'v4 Level', 'Driver', 'Timeline']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
risks = [
    ('Systemic ESG', 'CRITICAL', 'CRITICAL', 'UN inquiry', 'Immediate'),
    ('Legal', 'CRITICAL', 'CRITICAL', 'MPF findings', 'Immediate'),
    ('Reputational', 'CRITICAL', 'HIGH', 'Norway PM intervention', '30 days'),
    ('UN/Sovereign', 'N/A', 'CRITICAL', 'UN inquiry opened', '90 days'),
    ('Financial', 'HIGH', 'HIGH', 'Moody\'s review', '60 days'),
    ('Operational', 'HIGH', 'MEDIUM', 'Fosen compliance plan', '90 days'),
    ('Strategic', 'HIGH', 'HIGH', 'Pattern confirmed', 'Ongoing'),
]
for i, (r, v3, v4, d, t) in enumerate(risks, 1):
    table.rows[i].cells[0].text = r
    table.rows[i].cells[1].text = v3
    table.rows[i].cells[2].text = v4
    table.rows[i].cells[3].text = d
    table.rows[i].cells[4].text = t
doc.add_page_break()

# SECTION 9
doc.add_heading('9. GAP ANALYSIS', level=1)
doc.add_heading('What May Have Been Missed', level=2)
doc.add_paragraph("• UN inquiry internal documentation", style='List Bullet')
doc.add_paragraph("• Norway-Sami negotiations behind closed doors", style='List Bullet')
doc.add_paragraph("• Statkraft Board deliberations on Fosen compliance", style='List Bullet')
doc.add_paragraph("• MPF full report (preliminary only)", style='List Bullet')
doc.add_heading('v5 Follow-Up', level=2)
doc.add_paragraph("• Monitor Fosen compliance implementation weekly", style='List Bullet')
doc.add_paragraph("• Track UN inquiry preliminary findings", style='List Bullet')
doc.add_paragraph("• Watch Moody\'s ESG rating decision", style='List Bullet')
doc.add_page_break()

# SECTION 10
doc.add_heading('10. STRATEGIC RECOMMENDATIONS', level=1)
doc.add_heading('Immediate (Next 7 Days)', level=2)
doc.add_paragraph("1. FOSEN: Publicly welcome PM compliance announcement; accelerate timeline", style='List Number')
doc.add_paragraph("2. UN: Cooperate fully with Special Rapporteur inquiry", style='List Number')
doc.add_paragraph("3. BRAZIL: Respond to MPF findings with concrete remediation plan", style='List Number')
doc.add_paragraph("4. MOODY'S: Proactive engagement on ESG rating review", style='List Number')
doc.add_heading('Short-Term (30 Days)', level=2)
doc.add_paragraph("1. FOSEN: Publish detailed compliance roadmap with milestones", style='List Number')
doc.add_paragraph("2. BRAZIL: Federal-level indigenous consultation protocol", style='List Number')
doc.add_paragraph("3. GLOBAL: Company-wide FPIC implementation framework", style='List Number')
doc.add_heading('Long-Term (90 Days)', level=2)
doc.add_paragraph("1. Turn UN inquiry into transparency showcase", style='List Number')
doc.add_paragraph("2. Best-practice indigenous rights leadership", style='List Number')
doc.add_paragraph("3. Sector-wide standards proposal", style='List Number')

# APPENDIX
doc.add_page_break()
doc.add_heading('APPENDIX: v4 METHODOLOGY', level=1)
doc.add_heading('Version Building Protocol', level=2)
doc.add_paragraph("v4 preserves ALL v3 content and adds new week developments:")
doc.add_paragraph("• v3 baseline: Fosen, Los Lagos, Santa Eugênia, Devoll analysis")
doc.add_paragraph("• v4 additions: Norway PM announcement, UN inquiry, MPF findings, Moody's review")
doc.add_paragraph("• Trend analysis: 4-week progression (v1→v2→v3→v4)")
doc.add_paragraph("• Risk evolution: Tracks how each project risk changes week-over-week")

# Save
output = '/Users/jonathonmilne/.openclaw/workspace/Statkraft_Media_Monitoring_v4_Enterprise_COMPLETE.docx'
doc.save(output)
print(f"COMPLETE v4 with all 10 sections: {output}")
