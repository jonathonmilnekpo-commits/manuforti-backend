#!/usr/bin/env python3
"""Convert all Crab markdown files to Word documents"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

def md_to_docx(md_file, output_file):
    """Convert markdown file to Word document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split into lines
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 33, 71)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(43, 108, 176)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Tables (simplified - just convert to text)
        elif line.startswith('|') and '---' not in line:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                p = doc.add_paragraph(' | '.join(cells))
                p.style = 'Intense Quote'
        
        # Code blocks
        elif line.startswith('```'):
            continue  # Skip code block markers
        elif line.startswith('    ') or line.startswith('\t'):
            p = doc.add_paragraph(line.strip())
            p.style = 'Quote'
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(line[line.find('.')+1:].strip(), style='List Number')
        
        # Bold text
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        
        # Regular paragraph
        else:
            # Handle inline bold and italic
            p = doc.add_paragraph()
            
            # Split by bold markers
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith('*') and part.endswith('*'):
                    p.add_run(part[1:-1]).italic = True
                else:
                    p.add_run(part)
    
    # Save
    doc.save(output_file)
    print(f"✓ Converted: {md_file} → {output_file}")

# Convert all files
files_to_convert = [
    ('/Users/jonathonmilne/.openclaw/workspace/crab/SOUL.md', '/Users/jonathonmilne/.openclaw/workspace/crab/Crab_SOUL.docx'),
    ('/Users/jonathonmilne/.openclaw/workspace/crab/business_plan.md', '/Users/jonathonmilne/.openclaw/workspace/crab/Crab_Business_Plan_Markdown.docx'),
    ('/Users/jonathonmilne/.openclaw/workspace/crab/content_calendar_week1.md', '/Users/jonathonmilne/.openclaw/workspace/crab/Crab_Content_Calendar_Week1.docx'),
    ('/Users/jonathonmilne/.openclaw/workspace/crab/social_media_bios.md', '/Users/jonathonmilne/.openclaw/workspace/crab/Crab_Social_Media_Bios.docx'),
    ('/Users/jonathonmilne/.openclaw/workspace/crab/x_launch_posts.md', '/Users/jonathonmilne/.openclaw/workspace/crab/Crab_X_Launch_Posts.docx'),
    ('/Users/jonathonmilne/.openclaw/workspace/crab/scripts/week1_monday_morning_routine.md', '/Users/jonathonmilne/.openclaw/workspace/crab/Crab_Video_Script_Week1_Monday.docx'),
]

for md_file, docx_file in files_to_convert:
    if Path(md_file).exists():
        md_to_docx(md_file, docx_file)
    else:
        print(f"⚠ File not found: {md_file}")

print("\n✓ All conversions complete!")
