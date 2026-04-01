#!/usr/bin/env python3
"""Generate Crab Business Plan Word Document"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    if level == 1:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 33, 71)
            run.font.size = Pt(18)
            run.font.bold = True
    elif level == 2:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(43, 108, 176)
            run.font.size = Pt(14)
    return heading

doc = Document()

# Title
title = doc.add_heading('Crab Business Plan', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 33, 71)
    run.font.size = Pt(24)
    run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Faceless Content Empire — 6-Month Roadmap')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(113, 128, 150)
run.font.italic = True

doc.add_paragraph()

# Executive Summary
add_heading_custom(doc, 'Executive Summary', 1)
doc.add_paragraph('Crab operates as a distributed agent system creating faceless content across YouTube, Instagram, and X. Content themes focus on AI automation, productivity workflows, and build in public OpenClaw demonstrations.')

doc.add_paragraph('Revenue comes from:')
doc.add_paragraph('OpenClaw Setup Services (primary) — €500-5,000 per client', style='List Bullet')
doc.add_paragraph('Digital Products — €50-200 per sale', style='List Bullet')
doc.add_paragraph('Affiliate Revenue — passive', style='List Bullet')
doc.add_paragraph('Sponsorships (Month 4+) — €500-2,000 per deal', style='List Bullet')

# Revenue Targets
add_heading_custom(doc, '6-Month Revenue Targets', 1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
headers = ['Phase', 'Months', 'Revenue Target']
for i, h in enumerate(headers):
    hdr[i].text = h
    set_cell_shading(hdr[i], '002147')
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.bold = True

phases = [
    ('Foundation', '1-2', '€500-2,000'),
    ('Growth', '3-4', '€2,000-5,000'),
    ('Scale', '5-6', '€5,000-15,000/month'),
]

for phase, months, revenue in phases:
    row = table.add_row().cells
    row[0].text = phase
    row[1].text = months
    row[2].text = revenue

doc.add_paragraph()

# Phase 1
add_heading_custom(doc, 'Phase 1: Foundation (Months 1-2)', 1)
doc.add_paragraph('Target: €500-2,000 revenue')

add_heading_custom(doc, 'Goals', 2)
doc.add_paragraph('Launch 3 content channels (YouTube, Instagram, X)', style='List Bullet')
doc.add_paragraph('Publish 20+ pieces of content', style='List Bullet')
doc.add_paragraph('Build email list to 200 subscribers', style='List Bullet')
doc.add_paragraph('Close first 2-3 OpenClaw setup clients', style='List Bullet')

add_heading_custom(doc, 'Content Strategy', 2)

doc.add_paragraph('YouTube (Long-form):', style='List Bullet').bold = True
doc.add_paragraph('Format: Screen recordings + AI voiceover', style='List Bullet 2')
doc.add_paragraph('Frequency: 2 videos/week', style='List Bullet 2')
doc.add_paragraph('Length: 8-15 minutes', style='List Bullet 2')

doc.add_paragraph('Instagram (Short-form):', style='List Bullet').bold = True
doc.add_paragraph('Format: Reels with captions + trending audio', style='List Bullet 2')
doc.add_paragraph('Frequency: 1 reel/day', style='List Bullet 2')

doc.add_paragraph('X/Twitter:', style='List Bullet').bold = True
doc.add_paragraph('Format: Text threads + video clips', style='List Bullet 2')
doc.add_paragraph('Frequency: 2 threads/week + 3-5 daily tweets', style='List Bullet 2')

# Phase 2
add_heading_custom(doc, 'Phase 2: Growth (Months 3-4)', 1)
doc.add_paragraph('Target: €2,000-5,000 revenue')

add_heading_custom(doc, 'Goals', 2)
doc.add_paragraph('100+ YouTube subscribers, 1,000+ Instagram followers, 500+ X followers', style='List Bullet')
doc.add_paragraph('Email list: 500 subscribers', style='List Bullet')
doc.add_paragraph('5-8 OpenClaw setup clients', style='List Bullet')
doc.add_paragraph('Launch first digital product', style='List Bullet')

# Phase 3
add_heading_custom(doc, 'Phase 3: Scale (Months 5-6)', 1)
doc.add_paragraph('Target: €5,000-15,000/month revenue')

add_heading_custom(doc, 'Goals', 2)
doc.add_paragraph('500+ YouTube subscribers, 5,000+ Instagram followers, 2,000+ X followers', style='List Bullet')
doc.add_paragraph('Email list: 1,000+ subscribers', style='List Bullet')
doc.add_paragraph('10-15 clients/month or €5,000+ in mixed revenue', style='List Bullet')
doc.add_paragraph('First sponsorship deals', style='List Bullet')

# Sub-Agents
add_heading_custom(doc, 'Sub-Agent Architecture', 1)
doc.add_paragraph('Crab operates as an orchestrator, delegating to specialized sub-agents:')

agents = [
    ('Strategist', 'Content Strategy', 'Trend research, calendar planning, competitor analysis'),
    ('Scribe', 'Script Writer', 'Video scripts, thread copy, captions, email sequences'),
    ('Producer', 'Video Production', 'Editing, thumbnails, captions, multi-format exports'),
    ('Poster', 'Social Media Manager', 'Publishing, engagement, metrics tracking'),
    ('Qualifier', 'Lead Qualification', 'Inbound handling, lead scoring, call scheduling'),
    ('Builder', 'Client Delivery', 'OpenClaw setups, configuration, documentation'),
    ('Maker', 'Product Developer', 'Templates, digital products, landing pages'),
]

for name, role, tasks in agents:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{name} — {role}').bold = True
    doc.add_paragraph(tasks, style='List Bullet 2')

# Service Pricing
add_heading_custom(doc, 'Service Pricing', 1)

pricing_table = doc.add_table(rows=1, cols=4)
pricing_table.style = 'Table Grid'
hdr = pricing_table.rows[0].cells
headers = ['Tier', 'Price', 'Timeline', 'Best For']
for i, h in enumerate(headers):
    hdr[i].text = h
    set_cell_shading(hdr[i], '2B6CB0')
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.bold = True

pricing_data = [
    ('Starter', '€500-1,000', '2-3 hours', 'Individuals'),
    ('Professional', '€2,000-5,000', '1-2 days', 'Small teams'),
    ('Enterprise', '€10,000-50,000', '1-4 weeks', 'Organizations'),
]

for tier, price, timeline, best_for in pricing_data:
    row = pricing_table.add_row().cells
    row[0].text = tier
    row[1].text = price
    row[2].text = timeline
    row[3].text = best_for

doc.add_paragraph()

# KPIs
add_heading_custom(doc, 'Key Performance Indicators', 1)

add_heading_custom(doc, 'Content Metrics', 2)
doc.add_paragraph('Views per video: 500+ by Month 3, 2,000+ by Month 6', style='List Bullet')
doc.add_paragraph('Follower growth: 10% month-over-month', style='List Bullet')
doc.add_paragraph('Engagement rate: 5%+ on Instagram, 3%+ on X', style='List Bullet')
doc.add_paragraph('Email list: 200 → 1,000 subscribers', style='List Bullet')

add_heading_custom(doc, 'Revenue Metrics', 2)
doc.add_paragraph('Monthly recurring revenue (MRR)', style='List Bullet')
doc.add_paragraph('Client acquisition cost (CAC)', style='List Bullet')
doc.add_paragraph('Average deal size', style='List Bullet')
doc.add_paragraph('Conversion rate (lead to client)', style='List Bullet')

# Footer
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Crab — Build in public. Automate everything. Ship daily.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(113, 128, 150)
run.font.italic = True

# Save
output_path = '/Users/jonathonmilne/.openclaw/workspace/crab/Crab_Business_Plan.docx'
doc.save(output_path)
print(f'✓ Business plan saved to: {output_path}')
