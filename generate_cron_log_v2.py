#!/usr/bin/env python3
"""
Generate Venture Cron Job Log v2 Word Document
Combines information from all sources into a single DOCX
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_custom(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 33, 71)  # Navy
            run.font.size = Pt(18)
            run.font.bold = True
    elif level == 2:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(43, 108, 176)  # Steel Blue
            run.font.size = Pt(14)
    return heading

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Venture Cron Job Log v2', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 33, 71)
        run.font.size = Pt(24)
        run.font.bold = True
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Manu Forti Intelligence — Venture Agent Nightly Operations')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(113, 128, 150)
    run.font.italic = True
    
    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Last Updated: March 13, 2026 | Agent: Venture (📊) | Reports to: Aiden (🤝)')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(113, 128, 150)
    
    doc.add_paragraph()  # Spacer
    
    # AGENT CONFIGURATION
    add_heading_custom(doc, 'AGENT CONFIGURATION', 1)
    
    doc.add_paragraph('Name: Venture')
    doc.add_paragraph('Role: Manu Forti Intelligence venture development agent')
    doc.add_paragraph('Mission: Develop the Manu Forti Intelligence venture arm through consistent nightly execution on Product 1, website, and go-to-market')
    doc.add_paragraph('Session Schedule: 02:00–04:00 GMT nightly (2-hour window)')
    
    add_heading_custom(doc, 'Critical: Continuity Protocol', 2)
    
    doc.add_paragraph('BEFORE starting work, MUST read:', style='List Bullet')
    doc.add_paragraph('AGENT_VENTURE.md — configuration and long-term goals', style='List Bullet 2')
    doc.add_paragraph('memory/YYYY-MM-DD.md — last night\'s work', style='List Bullet 2')
    doc.add_paragraph('MEMORY.md — venture context and decisions', style='List Bullet 2')
    doc.add_paragraph('Any project files you were working on', style='List Bullet 2')
    
    doc.add_paragraph('AT END of session (03:50 GMT), MUST:', style='List Bullet')
    doc.add_paragraph('Write detailed summary to memory/YYYY-MM-DD.md', style='List Bullet 2')
    doc.add_paragraph('Update AGENT_VENTURE.md with new goals/learnings', style='List Bullet 2')
    doc.add_paragraph('Update MEMORY.md if significant decisions made', style='List Bullet 2')
    doc.add_paragraph('EIRIK REVIEW: Run skeptic review and log to learnings/EIRIK_LOG.md', style='List Bullet 2')
    
    # PRODUCT PORTFOLIO
    add_heading_custom(doc, 'PRODUCT PORTFOLIO & ROADMAP', 1)
    add_heading_custom(doc, 'Current Products (LOCKED — March 13, 2026)', 2)
    
    # Create product table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Product', 'Tier', 'Price', 'Buyer', 'Status']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], '002147')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    products = [
        ('Product 1', 'Standard', '€249', 'Category Manager', '✅ Ready'),
        ('Product 1', 'Premium', '€349', 'Category Manager', '✅ Ready'),
        ('Product 1', '3-Supplier Bundle', '€699', 'Category Manager', '✅ Ready'),
        ('Media Monitoring', 'Monitor', '€35/month', 'Category Manager', '✅ Ready'),
        ('Media Monitoring', 'Alert', '€105/month', 'Category Manager', '✅ Ready'),
        ('Category Strategy', 'Full', '€2,499', 'Procurement Director', '🔄 In Build'),
        ('Category Strategy', 'Bundle', '€3,999', 'Procurement Director', '🔄 In Build'),
        ('Academy', 'Practitioner', '€2,000–3,500', 'Category Manager', '✅ Ready'),
        ('Academy', 'Executive', '€8,000–12,000', 'CPO / VP', '✅ Ready'),
    ]
    
    for product, tier, price, buyer, status in products:
        row_cells = table.add_row().cells
        row_cells[0].text = product
        row_cells[1].text = tier
        row_cells[2].text = price
        row_cells[3].text = buyer
        row_cells[4].text = status
    
    doc.add_paragraph()
    
    # Phase Roadmap
    add_heading_custom(doc, 'Phase Roadmap', 2)
    
    doc.add_paragraph('Phase 1: Foundation (Q1-Q2 2026)', style='List Bullet').bold = True
    doc.add_paragraph('Product 1 MVP complete (9-slide reports, 8+ samples) ✅', style='List Bullet 2')
    doc.add_paragraph('Website v1 with payment flow 🔄', style='List Bullet 2')
    doc.add_paragraph('First 10 paying Product 1 customers ⏳', style='List Bullet 2')
    doc.add_paragraph('Backend API and order management system ⏳', style='List Bullet 2')
    
    doc.add_paragraph('Phase 2: Scale (Q3-Q4 2026)', style='List Bullet').bold = True
    doc.add_paragraph('Launch Product 2: Procurement Leadership Academy pilot', style='List Bullet 2')
    doc.add_paragraph('Customer acquisition: 50+ Product 1 customers', style='List Bullet 2')
    
    doc.add_paragraph('Phase 3: Platform (2027)', style='List Bullet').bold = True
    doc.add_paragraph('Integrated platform: Reports + Academy + Community', style='List Bullet 2')
    doc.add_paragraph('International expansion (EU, US markets)', style='List Bullet 2')
    
    # CRON JOB TASK HISTORY
    add_heading_custom(doc, 'CRON JOB TASK HISTORY', 1)
    
    # March 6
    add_heading_custom(doc, 'March 6, 2026 — Mission Brief', 2)
    
    doc.add_paragraph('TASK 1: Backend API Development', style='List Bullet').bold = True
    doc.add_paragraph('Outcome: Functional backend to make website production-ready', style='List Bullet 2')
    doc.add_paragraph('Set up Node.js/Express or Python/FastAPI server', style='List Bullet 3')
    doc.add_paragraph('Create PostgreSQL database schema for orders', style='List Bullet 3')
    doc.add_paragraph('Implement POST /api/orders endpoint', style='List Bullet 3')
    doc.add_paragraph('Stripe + Vipps payment integration', style='List Bullet 3')
    doc.add_paragraph('SendGrid for transactional emails', style='List Bullet 3')
    
    doc.add_paragraph('TASK 2: Website Redesign — Stripe-Inspired', style='List Bullet').bold = True
    doc.add_paragraph('Hero Section with dark navy gradient', style='List Bullet 3')
    doc.add_paragraph('Trust Signals + Features Grid', style='List Bullet 3')
    doc.add_paragraph('How It Works timeline', style='List Bullet 3')
    doc.add_paragraph('Pricing Section + Testimonials', style='List Bullet 3')
    
    doc.add_paragraph('TASK 3: Additional Sample Reports', style='List Bullet').bold = True
    doc.add_paragraph('Target: 1-2 more Product 1 reports (ABB, Siemens Energy)', style='List Bullet 2')
    
    # March 12
    add_heading_custom(doc, 'March 12, 2026 — Session Summary', 2)
    
    doc.add_paragraph('Focus: Website Product 2 Integration — Academy Landing Page', style='List Bullet')
    doc.add_paragraph('Academy Landing Page created: manuforti-website/academy.html ✅', style='List Bullet 2')
    doc.add_paragraph('Updated homepage with Academy nav and Product 2 pricing ✅', style='List Bullet 2')
    
    doc.add_paragraph('Key Decisions:', style='List Bullet').bold = True
    doc.add_paragraph('Mailto form for MVP (pilot recruitment)', style='List Bullet 2')
    doc.add_paragraph('Pilot-first positioning (€5,000 as primary CTA)', style='List Bullet 2')
    doc.add_paragraph('Cross-link strategy: Reports → Academy upsell', style='List Bullet 2')
    
    # March 13
    add_heading_custom(doc, 'March 13, 2026 — Session Summary', 2)
    
    doc.add_paragraph('Focus: Category Strategy Product Build + Methodology', style='List Bullet')
    doc.add_paragraph('Category Strategy Methodology Document (9 sections) ✅', style='List Bullet 2')
    doc.add_paragraph('Hard approval gate system implemented', style='List Bullet 2')
    doc.add_paragraph('Eirik mandatory review at end of every cron', style='List Bullet 2')
    
    doc.add_paragraph('Tonight\'s Task Order (6 Tasks):', style='List Bullet').bold = True
    doc.add_paragraph('Read methodology → write METHODOLOGY_REVIEW.md → WAIT FOR APPROVAL ⛔', style='List Bullet 2')
    doc.add_paragraph('Build Excel Template (BLOCKED until approval)', style='List Bullet 2')
    doc.add_paragraph('Build MCDM Calculator Script (BLOCKED)', style='List Bullet 2')
    doc.add_paragraph('Build Word Template (BLOCKED)', style='List Bullet 2')
    doc.add_paragraph('Update order.html (BLOCKED)', style='List Bullet 2')
    doc.add_paragraph('Update AGENT_VENTURE.md + Eirik review (BLOCKED)', style='List Bullet 2')
    
    # EIRIK'S PROTOCOL
    add_heading_custom(doc, 'EIRIK\'S DAILY REVIEW PROTOCOL', 1)
    
    doc.add_paragraph('Eirik (🔴) — Skeptic Agent')
    doc.add_paragraph('Role: Devil\'s advocate / critical challenger')
    doc.add_paragraph('Runs: End of every Venture cron job — mandatory')
    doc.add_paragraph('Log: learnings/EIRIK_LOG.md')
    
    add_heading_custom(doc, 'Review Questions (Answer All 5)', 2)
    
    doc.add_paragraph('Does this move us towards a paying customer? YES / NO / PARTIALLY', style='List Number')
    doc.add_paragraph('Is this the highest-priority thing to build?', style='List Number')
    doc.add_paragraph('Quality check: Is what was built good enough to sell?', style='List Number')
    doc.add_paragraph('Risk flag: Any new risks introduced today?', style='List Number')
    doc.add_paragraph('Tomorrow\'s priority (Eirik\'s recommendation)?', style='List Number')
    
    add_heading_custom(doc, 'Scoring', 2)
    doc.add_paragraph('🟢 Good — Moves business forward, right priority')
    doc.add_paragraph('🟡 Acceptable — Could be better but not wrong direction')
    doc.add_paragraph('🔴 Wrong priority — Building products without moving to paying customers')
    
    # Quote
    quote = doc.add_paragraph()
    quote_run = quote.add_run('"The single biggest risk going forward is not the SVP decision — it\'s the temptation to keep building products instead of preparing to sell them. Every cron job that adds a new product to the roadmap without deploying the backend is moving in the wrong direction. Build the infrastructure. When the SVP decision lands, be ready to sell in 48 hours."')
    quote_run.font.italic = True
    quote_run.font.color.rgb = RGBColor(113, 128, 150)
    
    # STRATEGIC CONSTRAINTS
    add_heading_custom(doc, 'STRATEGIC CONSTRAINTS (LOCKED)', 1)
    
    add_heading_custom(doc, 'SVP Recruitment Constraint', 2)
    doc.add_paragraph('Status: Jonathon interviewing for SVP Procurement at Statkraft')
    doc.add_paragraph('Impact: No public customer acquisition until SVP outcome known')
    doc.add_paragraph('Approach: Quiet build only — infrastructure, products, sales materials')
    doc.add_paragraph('Trigger: SVP decision → deploy backend → activate Stripe → first paying customers within 48hrs')
    
    # PRICING
    add_heading_custom(doc, 'PRICING VALIDATION (Eirik-Approved)', 1)
    
    pricing_table = doc.add_table(rows=1, cols=5)
    pricing_table.style = 'Table Grid'
    hdr = pricing_table.rows[0].cells
    headers = ['Product', 'Time', 'Price', 'Margin', 'Status']
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], '2B6CB0')
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True
    
    pricing_data = [
        ('Product 1 Standard', '~45 min', '€249', '~40%', '✅'),
        ('Product 1 Premium', '~1 hr', '€349', '~43%', '✅'),
        ('Category Strategy', '~4-5 hrs', '€2,499', '~65%', '✅ Best'),
        ('Category Bundle', '~6-7 hrs', '€3,999', '~65%', '✅'),
        ('Negotiation Brief', '~1.5 hrs', '€399', '~25%', '⚠️ Reprice to €499'),
    ]
    
    for prod, time, price, margin, status in pricing_data:
        row = pricing_table.add_row().cells
        row[0].text = prod
        row[1].text = time
        row[2].text = price
        row[3].text = margin
        row[4].text = status
    
    # FILE LOCATIONS
    add_heading_custom(doc, 'FILE LOCATIONS', 1)
    
    add_heading_custom(doc, 'Agent Configuration', 2)
    doc.add_paragraph('AGENT_VENTURE.md — Main configuration and goals')
    doc.add_paragraph('memory/YYYY-MM-DD.md — Daily session logs')
    doc.add_paragraph('learnings/EIRIK_LOG.md — Daily skeptic reviews')
    
    add_heading_custom(doc, 'Product 1', 2)
    doc.add_paragraph('Template: product1_v15_canonical_template.json')
    doc.add_paragraph('Generator: product1_generator_bulletproof.py')
    doc.add_paragraph('Quality Gate: product1_quality_gate.py')
    doc.add_paragraph('Reference: Boskalis_Product1_v15_Final.pptx')
    
    add_heading_custom(doc, 'Category Strategy (In Build)', 2)
    doc.add_paragraph('Methodology: CategoryStrategy_Methodology_v1.docx')
    doc.add_paragraph('Process: CATEGORY_STRATEGY_PROCESS.md')
    doc.add_paragraph('Template (pending): category-strategy-process/CATEGORY_STRATEGY_TEMPLATE.xlsx')
    doc.add_paragraph('MCDM Script (pending): category-strategy-process/MCDM_CALCULATOR.py')
    
    add_heading_custom(doc, 'Website', 2)
    doc.add_paragraph('Home: manuforti-website/index.html')
    doc.add_paragraph('Academy: manuforti-website/academy.html')
    doc.add_paragraph('Order: manuforti-website/order.html')
    doc.add_paragraph('Category Strategy: manuforti-website/category-strategy.html')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Document compiled from: AGENT_VENTURE.md, memory/2026-03-13.md, learnings/EIRIK_LOG.md, venture_tonight_task.txt')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(113, 128, 150)
    run.font.italic = True
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = footer2.add_run('Last compiled: March 13, 2026')
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(113, 128, 150)
    
    # Save
    output_path = '/Users/jonathonmilne/.openclaw/workspace/Venture_Cron_Job_Log_v2.docx'
    doc.save(output_path)
    print(f"✓ Document saved to: {output_path}")

if __name__ == '__main__':
    main()
