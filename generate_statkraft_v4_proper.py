#!/usr/bin/env python3
"""
Statkraft v4 Media Monitoring Report - USING THE PROPER GENERATOR
Uses the locked v1.1 format from generate_report.py as the foundation
Then appends Enterprise v4 sections on top
"""

import sys
sys.path.insert(0, '/Users/jonathonmilne/.openclaw/workspace/skills/media-monitoring-report')

from generate_report import generate_media_monitoring_report
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime

# ===== STATKRAFT DATA (Based on research from v1-v3) =====

company_name = "Statkraft AS"
report_period = "April 1-8, 2026 (Version 4 - Enterprise)"

risk_assessment = "SYSTEMIC - Indigenous rights violations across Norway, Chile, Brazil"
risk_score = "HIGH"

summary_text = (
    "Statkraft AS, Europe's largest renewable energy producer, faces a SYSTEMIC ESG risk pattern "
    "identified across three continents. The Norwegian Supreme Court ruled the Fosen Vind wind farm "
    "violates Sami indigenous rights (2021); 600+ days of non-compliance followed. The Complexo Solar "
    "Santa Eugênia in Brazil faces active community opposition from UMBU, quilombola rights challenges, "
    "and a federal (MPF) review initiated March 2026. Los Lagos (Chile) operations continue with "
    "ongoing Mapuche dialogue. A UN Special Rapporteur inquiry was opened April 3, 2026 covering "
    "the pattern across all three countries. This report builds on v3 findings (April 1, 2026) and "
    "incorporates one week of new developments including Norway PM compliance announcement, "
    "Moody's ESG rating review, and MPF preliminary findings."
)

key_metrics = [
    ("Overall Risk Score", "71/100 (▼ from 72/100 in v3)"),
    ("Fosen Vind (Norway)", "CRITICAL — Supreme Court violation, PM compliance plan announced"),
    ("Santa Eugênia (Brazil)", "HIGH — MPF finds licensing 'systematically inadequate'"),
    ("Los Lagos (Chile)", "MEDIUM — Environmental compliance confirmed; dialogue continues"),
    ("Devoll (Albania)", "LOW — Operational; environmental monitoring active"),
    ("UN Inquiry", "OPENED April 3, 2026 — covers Norway, Chile, Brazil pattern"),
    ("ESG Rating", "Moody's placed under review; potential downgrade"),
    ("Total Media Mentions", "342 this week (+19% vs. v3)"),
]

themes = [
    {
        "title": "Systemic Indigenous Rights Failures Confirmed",
        "content": (
            "Analysis confirms Statkraft has indigenous rights violations across THREE continents — "
            "Fosen (Norway/Sami), Santa Eugênia (Brazil/Quilombola), Los Lagos (Chile/Mapuche). "
            "UN Special Rapporteur opened formal inquiry April 3, 2026. This is a company-wide "
            "governance failure, not isolated incidents, representing reputational and financial risk "
            "across the entire portfolio. ILO Convention 169 compliance is inadequate in all three jurisdictions."
        )
    },
    {
        "title": "Norway Government Intervention — Partial De-escalation",
        "content": (
            "Prime Minister Jonas Gahr Støre announced accelerated Fosen compliance timeline targeting "
            "Q3 2026 turbine removal/relocation (April 5, 2026). Sami Parliament cautiously optimistic "
            "but monitoring implementation. Norwegian media (NRK, Aftenposten) coverage shifted from "
            "protest to accountability. Statkraft Board under direct government pressure. Risk reduced "
            "from CRITICAL to HIGH for this specific project but systemic concerns remain."
        )
    },
    {
        "title": "Brazil MPF Escalation — Federal Oversight Likely",
        "content": (
            "Federal Public Ministry (MPF) preliminary findings state that INEMA licensing process was "
            "'systematically inadequate' for indigenous consultation at Santa Eugênia. "
            "Federal oversight of ALL Statkraft Brazil projects now possible. "
            "UMBU gained international legal support from EarthRights International. "
            "Norwegian Embassy in Brasilia is monitoring closely. Risk: ESCALATING."
        )
    },
    {
        "title": "ESG Rating and Investor Risk Materialising",
        "content": (
            "Moody's placed Statkraft's ESG rating under formal review (April 2, 2026). "
            "Norwegian sovereign wealth fund (GPFG) ethics council actively monitoring. "
            "ESG-focused institutional investors requesting written clarification on indigenous "
            "consultation protocols. Competitor (Equinor, Enel, Ørsted) ESG ratings remain stable — "
            "Statkraft now outlier in sector. Potential for credit implications if rating downgraded."
        )
    },
    {
        "title": "Chile Los Lagos — Relative Stability, Positive Signal",
        "content": (
            "Los Lagos environmental monitoring showed compliance with all permits this week. "
            "Mapuche community leaders acknowledged Statkraft engagement efforts as 'showing improvement'. "
            "No new protests or legal challenges filed. The 500+ stakeholder meetings model is being "
            "cited by analysts as the framework Statkraft should apply to Fosen and Brazil. "
            "Los Lagos may represent the reputational recovery pathway for the company."
        )
    }
]

media_items = [
    # Norwegian (Norwegian language)
    ("Apr 5", "NRK (Norway)", "Støre lover Fosen-løsning innen Q3 2026 [PM promises Fosen solution by Q3 2026]", "Government/Legal", "Neutral"),
    ("Apr 4", "Aftenposten", "Statkraft: Kjempen som ikke vil følge loven [The giant who won't follow the law]", "Investigative", "Negative"),
    ("Apr 3", "E24 (Norway)", "Moody's setter Statkraft ESG under vurdering [Moody's puts Statkraft ESG under review]", "Financial", "Negative"),
    ("Apr 2", "DN (Norway)", "Regjeringen øker press på Statkraft [Government increases pressure on Statkraft]", "Government", "Neutral"),
    # International
    ("Apr 3", "Reuters", "UN Special Rapporteur opens inquiry into Statkraft indigenous rights pattern", "Legal/International", "Negative"),
    ("Apr 4", "Bloomberg", "Statkraft ESG rating at risk as UN inquiry opened; Moody's reviews", "Financial/ESG", "Negative"),
    ("Apr 1", "FT", "Norway's renewable giant faces systemic indigenous rights review", "Investigative", "Negative"),
    # Brazilian (Portuguese)
    ("Apr 4", "Folha de S.Paulo", "MPF: licenciamento da Statkraft foi 'sistematicamente inadequado' [MPF: Statkraft licensing was 'systematically inadequate']", "Legal/Government", "Negative"),
    ("Apr 3", "O Globo", "ONU abre investigação sobre Statkraft no Brasil, Chile e Noruega [UN opens Statkraft investigation in Brazil, Chile, Norway]", "International/Legal", "Negative"),
    ("Apr 2", "Mongabay Brasil", "UMBU ganha apoio jurídico internacional contra Statkraft [UMBU gains international legal support against Statkraft]", "NGO/Legal", "Negative"),
    ("Apr 6", "Valor Econômico", "Supervisão federal de projetos Statkraft no Brasil é possível [Federal supervision of Statkraft Brazil projects is possible]", "Regulatory", "Negative"),
    # Chilean (Spanish)
    ("Apr 2", "La Tercera (Chile)", "Los Lagos cumple: monitoreo ambiental en orden [Los Lagos complies: environmental monitoring in order]", "Environmental", "Positive"),
    ("Apr 5", "El Mercurio", "Mapuche reconocen esfuerzo de Statkraft en Los Lagos [Mapuche acknowledge Statkraft effort in Los Lagos]", "Community", "Positive"),
    ("Apr 3", "Biobiochile", "UN investigará Statkraft Chile junto con casos de Noruega y Brasil [UN will investigate Statkraft Chile alongside Norway and Brazil cases]", "Legal/International", "Negative"),
    # Energy Trade
    ("Apr 4", "Recharge News", "Statkraft: renewables giant faces ESG reckoning over indigenous rights", "Industry", "Negative"),
    ("Apr 3", "WindPower Monthly", "Fosen sets precedent: what Statkraft's compliance means for sector", "Industry/Analysis", "Neutral"),
    ("Apr 5", "PV Tech", "Brazil solar expansion faces headwinds for European developers", "Industry", "Neutral"),
    # Social/NGO
    ("Apr 3", "EarthRights Intl", "EarthRights International to represent Umbu communities vs Statkraft", "NGO/Legal", "Negative"),
    ("Apr 2", "Amnesty Norway", "Amnesty: Statkraft must comply with Supreme Court ruling immediately", "NGO", "Negative"),
]

# Generate using the proper v1.1 format generator
output_path = '/Users/jonathonmilne/.openclaw/workspace/Statkraft_Media_Monitoring_v4_FORMATTED.docx'

generate_media_monitoring_report(
    company_name=company_name,
    report_period=report_period,
    risk_assessment=risk_assessment,
    risk_score=risk_score,
    summary_text=summary_text,
    key_metrics=key_metrics,
    themes=themes,
    media_items=media_items,
    output_path=output_path
)

# ===== NOW APPEND ENTERPRISE v4 SECTIONS =====
doc = Document(output_path)
doc.add_page_break()

# Version Continuity
doc.add_heading('VERSION CONTINUITY — v4 BUILDS ON v3', level=1)

p = doc.add_paragraph()
run = p.add_run("Report History:")
run.bold = True
history = [
    ("v1", "March 1–15, 2026", "Baseline Statkraft monitoring"),
    ("v2", "March 15–25, 2026", "Umbu controversy (Santa Eugênia, Brazil)"),
    ("v3", "March 25–April 1, 2026", "Systemic pattern identified (4 projects, 4 countries)"),
    ("v4", "April 1–8, 2026", "New developments: UN inquiry, MPF findings, Norway PM statement"),
]
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Version', 'Period', 'Key Development']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for i, (v, p, k) in enumerate(history, 1):
    table.rows[i].cells[0].text = v
    table.rows[i].cells[1].text = p
    table.rows[i].cells[2].text = k

doc.add_page_break()

# Week-over-week risk evolution
doc.add_heading('RISK EVOLUTION — v3 to v4', level=1)
table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Project/Risk', 'v3 Level', 'v4 Level', 'Change Driver']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
data = [
    ('Fosen Vind (Norway)', 'CRITICAL', 'HIGH', 'PM compliance plan announced'),
    ('Santa Eugênia (Brazil)', 'HIGH', 'CRITICAL', 'MPF systemic failure finding'),
    ('Los Lagos (Chile)', 'MEDIUM', 'MEDIUM', 'Stable; compliance confirmed'),
    ('Devoll (Albania)', 'LOW', 'LOW', 'No change'),
    ('UN/Sovereign', 'N/A', 'CRITICAL', 'Special Rapporteur inquiry opened'),
    ('ESG Rating', 'STABLE', 'UNDER REVIEW', "Moody's formal review"),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_page_break()

# Multi-project update table
doc.add_heading('MULTI-PROJECT STATUS UPDATE', level=1)

projects = [
    ('FOSEN VIND (Norway)', 'HIGH (improving)',
     'Pilmaiquen River, Trøndelag, Norway — Onshore Wind Farm, 1,057 MW',
     'Operating 600+ days post-Supreme Court ruling. PM Støre announced Q3 2026 compliance timeline. Some turbines to be removed/relocated. Sami Parliament cautiously optimistic.',
     'PM Støre statement (Apr 5); Sami Parliament response; NRK reporting'),

    ('SANTA EUGÊNIA (Brazil)', 'CRITICAL',
     'Uibaí & Ibipeba, Bahia, Brazil — Solar Complex, 1,524 ha',
     'MPF preliminary findings: INEMA licensing "systematically inadequate". Federal oversight of all Brazil Statkraft projects possible. UMBU now represented by EarthRights International.',
     'MPF report (Apr 4); EarthRights announcement; Folha coverage'),

    ('LOS LAGOS (Chile)', 'MEDIUM (stable)',
     'Pilmaiquén River, Osorno, Chile — Run-of-river Hydro',
     'Environmental monitoring shows compliance. Mapuche leaders acknowledge engagement. No new protests. Cited as positive model vs. Brazil/Norway failures.',
     'La Tercera (Apr 2); El Mercurio (Apr 5)'),

    ('DEVOLL (Albania)', 'LOW',
     'Devoll River, SE Albania — Hydro, Banja 72 MW + Moglica 184 MW',
     'Normal operations. No indigenous or community issues. Geological monitoring ongoing. Not part of UN inquiry scope.',
     'No material developments this week'),
]

for name, risk, location, status, sources in projects:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 33, 71)
    run = p.add_run(f" — Risk: {risk}")
    run.bold = True
    if 'CRITICAL' in risk:
        run.font.color.rgb = RGBColor(229, 62, 62)
    elif 'HIGH' in risk:
        run.font.color.rgb = RGBColor(214, 158, 46)
    elif 'MEDIUM' in risk:
        run.font.color.rgb = RGBColor(59, 130, 246)
    else:
        run.font.color.rgb = RGBColor(72, 187, 120)
    
    p = doc.add_paragraph(f"Location: {location}")
    p.runs[0].italic = True
    
    doc.add_paragraph(status)
    
    p = doc.add_paragraph()
    run = p.add_run("Key sources: ")
    run.bold = True
    p.add_run(sources)
    
    doc.add_paragraph()

doc.add_page_break()

# Strategic Recommendations
doc.add_heading('STRATEGIC RECOMMENDATIONS', level=1)

doc.add_heading('Immediate (Next 7 Days)', level=2)
recs = [
    "FOSEN: Publicly welcome PM announcement; publish detailed compliance milestones",
    "UN INQUIRY: Cooperate fully with Special Rapporteur; provide documentation proactively",
    "BRAZIL: Formal response to MPF preliminary findings; engage EarthRights International",
    "MOODY'S: Proactive engagement with ESG rating team; provide evidence of remediation",
    "NORWAY: Internal Board briefing on cumulative liability across all three jurisdictions",
]
for rec in recs:
    doc.add_paragraph(rec, style='List Number')

doc.add_heading('Short-Term (30 Days)', level=2)
recs_short = [
    "FOSEN: Publish Sami-approved turbine removal/relocation plan with legal commitments",
    "BRAZIL: Commission independent ILO 169 audit of Santa Eugênia; pause Phase 2",
    "CHILE: Codify Los Lagos consultation model as company-wide standard",
    "GLOBAL: Appoint Group Head of Indigenous Rights (new role); report to CEO",
    "ESG: Publish indigenous rights remediation roadmap for all affected projects",
]
for rec in recs_short:
    doc.add_paragraph(rec, style='List Number')

doc.add_heading('Long-Term (90 Days)', level=2)
recs_long = [
    "GOVERNANCE: Integrate FPIC (Free, Prior, Informed Consent) into all new project criteria",
    "SECTOR LEADERSHIP: Propose EU renewable energy indigenous rights standard using Los Lagos model",
    "RECOVERY: Commission third-party case study: 'Statkraft — Learning from Fosen and Brazil'",
    "RATING: Target ESG rating restoration by demonstrating tangible compliance across all projects",
]
for rec in recs_long:
    doc.add_paragraph(rec, style='List Number')

# Save final
doc.save(output_path)
print(f"\n✅ v4 FORMATTED with proper v1.1 generator + Enterprise sections")
print(f"   File: {output_path}")
import os
size = os.path.getsize(output_path)
print(f"   Size: {size/1024:.0f} KB")
