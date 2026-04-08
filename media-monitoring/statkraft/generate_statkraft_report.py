#!/usr/bin/env python3
"""
Statkraft Media Monitoring Report Generator
============================================
Usage: python3 generate_statkraft_report.py

ALWAYS uses the locked generate_report.py skill — never writes formatting from scratch.
Updates version in statkraft_data.py, regenerates the full report with charts.
"""

import sys
import os
sys.path.insert(0, '/Users/jonathonmilne/.openclaw/workspace/skills/media-monitoring-report')

from generate_report import generate_media_monitoring_report
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ===== IMPORT DATA =====
sys.path.insert(0, '/Users/jonathonmilne/.openclaw/workspace/media-monitoring/statkraft')
import statkraft_data as data

OUTPUT_PATH = f'/Users/jonathonmilne/.openclaw/workspace/media-monitoring/statkraft/Statkraft_Media_Monitoring_{data.VERSION}_{data.REPORT_PERIOD[:10].replace(" ","_")}.docx'

print(f"Generating {data.VERSION} report...")
print(f"Period: {data.REPORT_PERIOD}")
print(f"Articles: {len(data.MEDIA_ITEMS)}")
print(f"Projects: {len(data.PROJECT_STATUS)}")

# ===== STEP 1: GENERATE USING PROPER FORMATTER =====
generate_media_monitoring_report(
    company_name=data.COMPANY_NAME,
    report_period=data.REPORT_PERIOD,
    risk_assessment=data.RISK_ASSESSMENT,
    risk_score=data.RISK_SCORE,
    summary_text=data.SUMMARY,
    key_metrics=data.KEY_METRICS,
    themes=data.THEMES,
    media_items=data.MEDIA_ITEMS,
    output_path=OUTPUT_PATH,
    social_data=data.SOCIAL_DATA,
)

# ===== STEP 2: APPEND ENTERPRISE SECTIONS =====
doc = Document(OUTPUT_PATH)

doc.add_page_break()

# ------- PROJECT-BY-PROJECT STATUS -------
h = doc.add_heading('PROJECT-BY-PROJECT STATUS', level=1)
h.runs[0].font.color.rgb = RGBColor(0, 33, 71)

table = doc.add_table(rows=len(data.PROJECT_STATUS) + 1, cols=4)
table.style = 'Light Grid Accent 1'

for i, hdr in enumerate(['Project', 'Risk', 'Status Summary', 'Latest Development']):
    table.rows[0].cells[i].text = hdr
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

risk_colors = {'CRITICAL': 'f8d7da', 'HIGH': 'fff3cd', 'MEDIUM': 'e2f0d9', 'LOW': 'd4edda'}

for i, (name, proj) in enumerate(data.PROJECT_STATUS.items(), 1):
    row = table.rows[i]
    row.cells[0].text = f"{name}\n{proj['type']}\n{proj['capacity']}"
    row.cells[1].text = proj['risk']
    row.cells[2].text = proj['status'][:300] + ('...' if len(proj['status']) > 300 else '')
    row.cells[3].text = proj['latest_development']
    # Colour code risk cell
    color = risk_colors.get(proj['risk'], 'ffffff')
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    row.cells[1]._tc.get_or_add_tcPr().append(shading)
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(9) if cell.paragraphs[0].runs else None

doc.add_page_break()

# ------- SEARCH PROTOCOL USED -------
h = doc.add_heading('RESEARCH PROTOCOL — SEARCHES CONDUCTED', level=1)
h.runs[0].font.color.rgb = RGBColor(0, 33, 71)

p = doc.add_paragraph()
p.add_run("Standard: Project names + company name in local language of each operating country\n").bold = True

for name, proj in data.PROJECT_STATUS.items():
    p = doc.add_paragraph()
    p.add_run(f"{name} ({proj['location']}): ").bold = True
    p.add_run(", ".join(proj['search_terms']))

doc.add_page_break()

# ------- STRATEGIC RECOMMENDATIONS -------
h = doc.add_heading('STRATEGIC RECOMMENDATIONS', level=1)
h.runs[0].font.color.rgb = RGBColor(0, 33, 71)

doc.add_heading('Immediate (Next 7 Days)', level=2)
for rec in [
    "FOSEN: Publish concrete Sør-Fosen engagement plan; address outstanding dispute with southern group",
    "BRAZIL: Respond publicly to MPF licensing review; demonstrate UMBU engagement",
    "CHILE: Escalate to CEO level — Mapuche have now met Norwegian Ambassador; this is sovereign-level",
    "ESG: Commission independent FPIC audit across all three indigenous rights cases",
    "COMMUNICATIONS: Prepare unified statement addressing the systemic pattern — do not treat as isolated",
]:
    doc.add_paragraph(rec, style='List Number')

doc.add_heading('Short-Term (30 Days)', level=2)
for rec in [
    "GOVERNANCE: Appoint Group Head of Indigenous Rights — report directly to CEO",
    "FOSEN: Codify Dec 2023 deal into binding legal instrument with full Sámi participation",
    "BRAZIL: Pause Sol de Brotas expansion pending UMBU community consent process",
    "CHILE: Commission independent ILO 169 audit of Los Lagos; publish results",
    "ESG REPORTING: Add FPIC compliance section to 2026 Annual Report",
]:
    doc.add_paragraph(rec, style='List Number')

doc.add_heading('Long-Term (90 Days)', level=2)
for rec in [
    "COMPANY-WIDE: Implement FPIC as mandatory gate before any project approval in new markets",
    "SECTOR LEADERSHIP: Use Los Lagos 500-meeting model as foundation for Statkraft FPIC standard",
    "REPUTATIONAL RECOVERY: Commission third-party case study — turn lessons learned into industry standard",
    "RATINGS: Engage S&P, Moody's and Sustainalytics on roadmap to ESG rating stabilisation",
]:
    doc.add_paragraph(rec, style='List Number')

# Save
doc.save(OUTPUT_PATH)
size = os.path.getsize(OUTPUT_PATH)
print(f"\n✅ {data.VERSION} report complete")
print(f"   File: {OUTPUT_PATH}")
print(f"   Size: {size/1024:.0f} KB")
print(f"   Articles: {len(data.MEDIA_ITEMS)}")
print(f"   Projects: {len(data.PROJECT_STATUS)}")
