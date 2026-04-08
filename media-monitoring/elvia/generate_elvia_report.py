#!/usr/bin/env python3
"""Elvia AS — Media Monitoring Report Generator"""

import sys, os
sys.path.insert(0, '/Users/jonathonmilne/.openclaw/workspace/skills/media-monitoring-report')
from generate_report import generate_media_monitoring_report

sys.path.insert(0, os.path.dirname(__file__))
import elvia_data as data

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

OUTPUT = f'/Users/jonathonmilne/.openclaw/workspace/media-monitoring/elvia/Elvia_Media_Monitoring_{data.VERSION}.docx'

print(f"Generating Elvia {data.VERSION} — {data.REPORT_PERIOD}")
print(f"  Articles: {len(data.MEDIA_ITEMS)} | Projects: {len(data.PROJECT_STATUS)}")

generate_media_monitoring_report(
    company_name=data.COMPANY_NAME,
    report_period=data.REPORT_PERIOD,
    risk_assessment=data.RISK_ASSESSMENT,
    risk_score=data.RISK_SCORE,
    summary_text=data.SUMMARY,
    key_metrics=data.KEY_METRICS,
    themes=data.THEMES,
    media_items=data.MEDIA_ITEMS,
    output_path=OUTPUT,
    social_data=data.SOCIAL_DATA,
)

doc = Document(OUTPUT)
doc.add_page_break()

# --- Project status table ---
h = doc.add_heading('PROJECT & ASSET STATUS', level=1)
h.runs[0].font.color.rgb = RGBColor(0, 33, 71)

risk_colors = {'CRITICAL':'f8d7da','HIGH':'fff3cd','MEDIUM':'fff3cd','LOW':'d4edda'}
table = doc.add_table(rows=len(data.PROJECT_STATUS)+1, cols=4)
table.style = 'Light Grid Accent 1'
for i, hdr in enumerate(['Asset/Project','Risk','Status','Latest Development']):
    table.rows[0].cells[i].text = hdr
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for i, (name, proj) in enumerate(data.PROJECT_STATUS.items(), 1):
    row = table.rows[i]
    row.cells[0].text = f"{name}\n{proj['type']}\n{proj['location']}"
    row.cells[1].text = proj['risk']
    row.cells[2].text = proj['status'][:300] + ('...' if len(proj['status'])>300 else '')
    row.cells[3].text = proj['latest_development']
    color = risk_colors.get(proj['risk'], 'ffffff')
    row.cells[1]._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_page_break()

# --- Search protocol ---
h = doc.add_heading('RESEARCH PROTOCOL — SEARCHES CONDUCTED', level=1)
h.runs[0].font.color.rgb = RGBColor(0, 33, 71)
p = doc.add_paragraph()
p.add_run("Elvia operates only in Norway. All searches conducted in Norwegian.\n").bold = True
p.add_run("Sources: NRK, VG, Aftenposten, Nettavisen, E24, Nationen, Gudbrandsdølen Dagningen, NTB, Elvia.no, Kraftsystemet.no, Forbrukerrådet, Heimdall Power, Siemens, NIB\n")
doc.add_paragraph()
for name, proj in data.PROJECT_STATUS.items():
    p = doc.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(", ".join(proj['search_terms']))

doc.add_page_break()

# --- Strategic Recommendations ---
h = doc.add_heading('STRATEGIC RECOMMENDATIONS', level=1)
h.runs[0].font.color.rgb = RGBColor(0, 33, 71)

doc.add_heading('Immediate (Next 7 Days)', level=2)
for rec in [
    "TARIFF COMMUNICATION: Proactive consumer communication on 2026 tariff structure — emphasise tax offset",
    "POLITICAL ENGAGEMENT: Brief key Storting members on investment rationale before further nettleie debate",
    "DIVIDEND NARRATIVE: Prepare response to ongoing criticism of NOK 535M dividend; link to owner investment mandate",
    "LIÅSEN: Publish project milestone update and consumer benefit communications",
]:
    doc.add_paragraph(rec, style='List Number')

doc.add_heading('Short-Term (30 Days)', level=2)
for rec in [
    "2030 ROADMAP: Publish transparent grid investment plan with year-by-year tariff outlook",
    "INNOVATION SHOWCASE: Commission case study on Heimdall/Siemens digital twin impact — position as European leader",
    "AMS STANDARDS: Engage Forbrukerrådet on HAN-port standardisation; address lock-in concerns proactively",
    "SECURITY COMMS: Publish updated emergency preparedness plan following beredskap increase",
]:
    doc.add_paragraph(rec, style='List Number')

doc.add_heading('Long-Term (90 Days)', level=2)
for rec in [
    "TARIFF REFORM: Work with NVE/RME on consumer-friendly tariff structure for 2027+",
    "DIGITAL LEADERSHIP: Apply for EU Horizon funding as reference DSO for smart grid innovation",
    "CUSTOMER TRUST: Annual customer satisfaction report with transparent investment vs. tariff linkage",
]:
    doc.add_paragraph(rec, style='List Number')

doc.save(OUTPUT)
size = os.path.getsize(OUTPUT)
print(f"\n✅ Report complete: {OUTPUT}")
print(f"   Size: {size//1024} KB")
