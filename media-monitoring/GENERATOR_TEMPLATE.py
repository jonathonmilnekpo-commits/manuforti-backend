#!/usr/bin/env python3
"""
Manu Forti Intelligence — Media Monitoring Generator Template
=============================================================
Copy to media-monitoring/<company>/generate_<company>_report.py
Replace COMPANY_MODULE with the company data module name (e.g. statkraft_data)

Usage: python3 generate_<company>_report.py
"""

import sys
import os

# ===== MANDATORY — DO NOT REMOVE =====
sys.path.insert(0, '/Users/jonathonmilne/.openclaw/workspace/skills/media-monitoring-report')
from generate_report import generate_media_monitoring_report  # ALWAYS use this — never write from scratch

# ===== IMPORT COMPANY DATA =====
sys.path.insert(0, os.path.dirname(__file__))
import COMPANY_MODULE as data  # replace COMPANY_MODULE with actual data file name

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

OUTPUT_PATH = os.path.join(
    os.path.dirname(__file__),
    f"{data.COMPANY_NAME.replace(' ', '_')}_{data.VERSION}.docx"
)

print(f"Generating {data.VERSION} — {data.REPORT_PERIOD}")

# ===== STEP 1: GENERATE BASE REPORT (cover, charts, themes, media table) =====
generate_media_monitoring_report(
    company_name=data.COMPANY_NAME,
    report_period=data.REPORT_PERIOD,
    risk_assessment=data.RISK_ASSESSMENT,
    risk_score=data.RISK_SCORE,
    summary_text=data.SUMMARY,
    key_metrics=data.KEY_METRICS,
    themes=data.THEMES,
    media_items=data.MEDIA_ITEMS,
    output_path=OUTPUT_PATH,
    social_data=data.SOCIAL_DATA,
)

# ===== STEP 2: APPEND ENTERPRISE SECTIONS =====
doc = Document(OUTPUT_PATH)

# --- Project-by-project status ---
doc.add_page_break()
h = doc.add_heading('PROJECT-BY-PROJECT STATUS', level=1)
h.runs[0].font.color.rgb = RGBColor(0, 33, 71)

risk_colors = {'CRITICAL': 'f8d7da', 'HIGH': 'fff3cd', 'MEDIUM': 'e2f0d9', 'LOW': 'd4edda'}
table = doc.add_table(rows=len(data.PROJECT_STATUS) + 1, cols=4)
table.style = 'Light Grid Accent 1'

for i, hdr in enumerate(['Project', 'Risk', 'Status Summary', 'Latest Development']):
    table.rows[0].cells[i].text = hdr
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for i, (name, proj) in enumerate(data.PROJECT_STATUS.items(), 1):
    row = table.rows[i]
    row.cells[0].text = f"{name}\n{proj['type']} | {proj['capacity']}"
    row.cells[1].text = proj['risk']
    row.cells[2].text = proj['status'][:280] + ('...' if len(proj['status']) > 280 else '')
    row.cells[3].text = proj['latest_development']
    color = risk_colors.get(proj['risk'], 'ffffff')
    row.cells[1]._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    )

# --- Search protocol ---
doc.add_page_break()
h = doc.add_heading('RESEARCH PROTOCOL — SEARCHES CONDUCTED', level=1)
h.runs[0].font.color.rgb = RGBColor(0, 33, 71)
p = doc.add_paragraph()
p.add_run("Standard: Project names + company name in local language of each operating country\n").bold = True
for name, proj in data.PROJECT_STATUS.items():
    p = doc.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(", ".join(proj['search_terms']))

# --- Strategic recommendations ---
doc.add_page_break()
h = doc.add_heading('STRATEGIC RECOMMENDATIONS', level=1)
h.runs[0].font.color.rgb = RGBColor(0, 33, 71)
# Add recommendations here specific to the company

doc.save(OUTPUT_PATH)
size = os.path.getsize(OUTPUT_PATH)
print(f"\n✅ Report complete: {OUTPUT_PATH} ({size//1024} KB)")
