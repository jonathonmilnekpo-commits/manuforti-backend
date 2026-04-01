from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_instruction_box(doc, text):
    """Add a grey instruction box"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E2E8F0")  # Light grey
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.style = 'Intense Quote'
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)  # Steel gray
    doc.add_paragraph()  # Spacing

def add_navy_heading(doc, text, level=1):
    """Add a navy-colored heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)  # Navy #002147
        run.font.bold = True
    return heading

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('CATEGORY STRATEGY', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)
    run.font.size = Pt(24)
    run.font.bold = True

subtitle = doc.add_paragraph('Strategic Sourcing & Procurement Analysis')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

doc.add_paragraph()

# ============================================
# CLIENT BRIEF (Front Matter)
# ============================================
add_navy_heading(doc, 'CLIENT BRIEF', level=1)

add_instruction_box(doc, "Complete this intake form at the start of the engagement. All fields inform the analysis.")

# Client Brief Table
brief_table = doc.add_table(rows=10, cols=2)
brief_table.style = 'Light Grid Accent 1'

brief_data = [
    ("Client Name", "[Enter client name]"),
    ("Category Name", "[Enter category, e.g., High Voltage Electrical Equipment]"),
    ("Annual Spend", "[€ amount]"),
    ("Currency", "[EUR/USD/GBP/NOK]"),
    ("Incumbent Suppliers", "[List 2-3 current suppliers]"),
    ("Key Pain Points", "[e.g., long lead times, single source risk]"),
    ("Strategic Priorities", "[e.g., cost reduction, resilience]"),
    ("Timeline Constraint", "[e.g., DG2 gate in 6 months]"),
    ("Evaluation Criteria", "[Cost 30%, Resilience 25%, Risk 20%, Strategic 15%, Ease 10%]"),
    ("Project Pipeline", "[e.g., 9 projects/yr: 2 large, 3 mid, 4 small]"),
]

for idx, (label, value) in enumerate(brief_data):
    row = brief_table.rows[idx]
    row.cells[0].text = label
    row.cells[1].text = value
    # Make label bold
    for run in row.cells[0].paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)

doc.add_page_break()

# ============================================
# SECTION 1: EXECUTIVE SUMMARY
# ============================================
add_navy_heading(doc, '1. EXECUTIVE SUMMARY', level=1)

add_instruction_box(doc, "One-paragraph problem statement. One-sentence financial summary (cost, return, ROI). One-sentence recommended strategy. Keep to half a page maximum.")

doc.add_paragraph("[PROBLEM STATEMENT: Describe the current situation and why change is needed. What pain points or opportunities drove this analysis?]")

doc.add_paragraph()

doc.add_paragraph("[FINANCIAL SUMMARY: This strategy delivers €X million in annual value at a cost of €Y thousand, representing a Z:1 ROI with payback in N months.]")

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("RECOMMENDED STRATEGY: ").bold = True
p.add_run("[State the recommended approach in one clear sentence, e.g., 'Implement a dual-source framework agreement with Siemens Energy and Hitachi Energy, supported by a 6-month transition roadmap.']")

doc.add_page_break()

# ============================================
# SECTION 2: STRATEGIC CONTEXT
# ============================================
add_navy_heading(doc, '2. STRATEGIC CONTEXT', level=1)

add_instruction_box(doc, "Position the category in the Kraljic matrix. Describe the current approach and why it's being questioned. Identify the key drivers for change.")

add_navy_heading(doc, '2.1 Kraljic Positioning', level=2)
doc.add_paragraph("[Classify this category: Strategic (high profit impact, high supply risk), Leverage (high profit impact, low supply risk), Bottleneck (low profit impact, high supply risk), or Routine (low profit impact, low supply risk). Explain the rationale.]")

add_navy_heading(doc, '2.2 Current Approach', level=2)
doc.add_paragraph("[Describe how this category is currently sourced. Single source? Multiple sources? Framework or spot? What has worked and what hasn't?]")

add_navy_heading(doc, '2.3 Drivers for Change', level=2)
doc.add_paragraph("[What has changed? Market conditions? Internal requirements? Risk events? Regulatory shifts? List 3-5 specific drivers.]")

doc.add_page_break()

# ============================================
# SECTION 3: MARKET ANALYSIS
# ============================================
add_navy_heading(doc, '3. MARKET ANALYSIS', level=1)

add_instruction_box(doc, "Provide a comprehensive view of the supplier landscape, market dynamics, pricing intelligence, and risk environment. Use tables where appropriate.")

add_navy_heading(doc, '3.1 Supplier Landscape', level=2)

# Supplier table
supplier_table = doc.add_table(rows=7, cols=5)
supplier_table.style = 'Light Grid Accent 1'
supplier_headers = ['Supplier', 'HQ', 'Revenue (€M)', 'Market Share', 'Notes']
for idx, header in enumerate(supplier_headers):
    cell = supplier_table.cell(0, idx)
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)

supplier_rows = [
    ["[Supplier 1]", "[Country]", "[Amount]", "[%]", "[Key notes]"],
    ["[Supplier 2]", "[Country]", "[Amount]", "[%]", "[Key notes]"],
    ["[Supplier 3]", "[Country]", "[Amount]", "[%]", "[Key notes]"],
    ["[Supplier 4]", "[Country]", "[Amount]", "[%]", "[Key notes]"],
    ["[Supplier 5]", "[Country]", "[Amount]", "[%]", "[Key notes]"],
    ["[Supplier 6]", "[Country]", "[Amount]", "[%]", "[Key notes]"],
]

for row_idx, row_data in enumerate(supplier_rows, 1):
    for col_idx, value in enumerate(row_data):
        supplier_table.cell(row_idx, col_idx).text = value

doc.add_paragraph()

add_navy_heading(doc, '3.2 Market Dynamics', level=2)
doc.add_paragraph("[Market size, growth rate (CAGR), supply/demand balance, lead time benchmarks, recent M&A activity.]")

add_navy_heading(doc, '3.3 Cost Driver Analysis', level=2)

# Cost driver table
driver_table = doc.add_table(rows=7, cols=5)
driver_table.style = 'Light Grid Accent 1'
driver_headers = ['Driver', '5yr Trend', 'Price Impact', 'Forward Projection', 'Risk Rating']
for idx, header in enumerate(driver_headers):
    cell = driver_table.cell(0, idx)
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)

driver_rows = [
    ["[Raw Material 1]", "[+/- X%]", "[X%]", "[Projection]", "[High/Med/Low]"],
    ["[Raw Material 2]", "[+/- X%]", "[X%]", "[Projection]", "[High/Med/Low]"],
    ["[Energy]", "[+/- X%]", "[X%]", "[Projection]", "[High/Med/Low]"],
    ["[Labour]", "[+/- X%]", "[X%]", "[Projection]", "[High/Med/Low]"],
    ["[Logistics]", "[+/- X%]", "[X%]", "[Projection]", "[High/Med/Low]"],
    ["[FX]", "[+/- X%]", "[X%]", "[Projection]", "[High/Med/Low]"],
]

for row_idx, row_data in enumerate(driver_rows, 1):
    for col_idx, value in enumerate(row_data):
        driver_table.cell(row_idx, col_idx).text = value

doc.add_paragraph()

add_navy_heading(doc, '3.4 Pricing Intelligence', level=2)
doc.add_paragraph("[List price benchmarks, negotiated price ranges, typical discount structures, payment term norms.]")

add_navy_heading(doc, '3.5 Risk Environment', level=2)
doc.add_paragraph("[Geopolitical risks, regulatory changes (CSRD/CSDDD), technology disruption, ESG hot spots.]")

doc.add_page_break()

# ============================================
# SECTION 4: PROCUREMENT PROCESS
# ============================================
add_navy_heading(doc, '4. PROCUREMENT PROCESS', level=1)

add_instruction_box(doc, "Map the current procurement process for this category. Identify decision gates, timelines, and key stakeholders.")

add_navy_heading(doc, '4.1 Current Process', level=2)
doc.add_paragraph("[Describe the end-to-end process from demand identification to contract award and management.]")

add_navy_heading(doc, '4.2 Decision Gates & Timelines', level=2)
doc.add_paragraph("[What are the key approval points? How long does each stage typically take? Where are the bottlenecks?]")

add_navy_heading(doc, '4.3 Stakeholder Map', level=2)
doc.add_paragraph("[Who are the key internal stakeholders? What are their interests and influence levels?]")

doc.add_page_break()

# ============================================
# SECTION 5: STRATEGIC OPTIONS
# ============================================
add_navy_heading(doc, '5. STRATEGIC OPTIONS', level=1)

add_instruction_box(doc, "Develop 4-6 mutually exclusive strategic options. Each option should have a clear description, mechanism, advantages, and risks.")

# Options table
options_table = doc.add_table(rows=6, cols=5)
options_table.style = 'Light Grid Accent 1'
options_headers = ['Option', 'Description', 'Mechanism', 'Advantages', 'Risks']
for idx, header in enumerate(options_headers):
    cell = options_table.cell(0, idx)
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)

options_data = [
    ["Status Quo", "Continue current approach", "No change", "Low disruption", "Missed value"],
    ["Option 1", "[Description]", "[How it works]", "[Benefits]", "[Drawbacks]"],
    ["Option 2", "[Description]", "[How it works]", "[Benefits]", "[Drawbacks]"],
    ["Option 3", "[Description]", "[How it works]", "[Benefits]", "[Drawbacks]"],
    ["Option 4", "[Description]", "[How it works]", "[Benefits]", "[Drawbacks]"],
]

for row_idx, row_data in enumerate(options_data, 1):
    for col_idx, value in enumerate(row_data):
        options_table.cell(row_idx, col_idx).text = value

doc.add_paragraph()

add_navy_heading(doc, '5.1 Narrative Description', level=2)
doc.add_paragraph("[Provide a paragraph for each option explaining the rationale, when it might be appropriate, and any real-world examples.]")

doc.add_page_break()

# ============================================
# SECTION 6: MCDM EVALUATION
# ============================================
add_navy_heading(doc, '6. MCDM EVALUATION', level=1)

add_instruction_box(doc, "Present the AHP weighting and TOPSIS scoring. Show the criteria, weights, option scores, and final ranking. Include interpretation of the results.")

add_navy_heading(doc, '6.1 AHP Criteria and Weights', level=2)

# AHP table
ahp_table = doc.add_table(rows=6, cols=3)
ahp_table.style = 'Light Grid Accent 1'
ahp_headers = ['Criterion', 'Weight', 'Rationale']
for idx, header in enumerate(ahp_headers):
    cell = ahp_table.cell(0, idx)
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)

ahp_data = [
    ["Cost Reduction", "30%", "TCO, unit price, lifecycle cost optimization"],
    ["Supply Resilience", "25%", "Lead time reduction, supply security, dual-source capability"],
    ["Risk Reduction", "20%", "Financial, geopolitical, ESG, concentration risk mitigation"],
    ["Strategic Alignment", "15%", "Partnership depth, innovation access, ESG goal alignment"],
    ["Implementation Ease", "10%", "Speed to implement, change management burden"],
]

for row_idx, row_data in enumerate(ahp_data, 1):
    for col_idx, value in enumerate(row_data):
        ahp_table.cell(row_idx, col_idx).text = value

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Consistency Ratio (CR): ").bold = True
p.add_run("[0.XX] - CR < 0.10 indicates acceptable consistency")

add_navy_heading(doc, '6.2 TOPSIS Scoring Matrix', level=2)

# TOPSIS table
topsis_table = doc.add_table(rows=7, cols=7)
topsis_table.style = 'Light Grid Accent 1'
topsis_headers = ['Option', 'Cost', 'Resilience', 'Risk', 'Strategic', 'Ease', 'TOPSIS Score']
for idx, header in enumerate(topsis_headers):
    cell = topsis_table.cell(0, idx)
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)

topsis_data = [
    ["[Option 1]", "[1-10]", "[1-10]", "[1-10]", "[1-10]", "[1-10]", "[0.XXX]"],
    ["[Option 2]", "[1-10]", "[1-10]", "[1-10]", "[1-10]", "[1-10]", "[0.XXX]"],
    ["[Option 3]", "[1-10]", "[1-10]", "[1-10]", "[1-10]", "[1-10]", "[0.XXX]"],
    ["[Option 4]", "[1-10]", "[1-10]", "[1-10]", "[1-10]", "[1-10]", "[0.XXX]"],
    ["[Option 5]", "[1-10]", "[1-10]", "[1-10]", "[1-10]", "[1-10]", "[0.XXX]"],
    ["[Option 6]", "[1-10]", "[1-10]", "[1-10]", "[1-10]", "[1-10]", "[0.XXX]"],
]

for row_idx, row_data in enumerate(topsis_data, 1):
    for col_idx, value in enumerate(row_data):
        topsis_table.cell(row_idx, col_idx).text = value

doc.add_paragraph()

add_navy_heading(doc, '6.3 Results and Commentary', level=2)
doc.add_paragraph("[Present the final ranking. Explain why the winner scored highest and what differentiates it from close competitors. Discuss sensitivity: would the ranking change if weights shifted?]")

doc.add_page_break()

# ============================================
# SECTION 7: RECOMMENDED STRATEGY
# ============================================
add_navy_heading(doc, '7. RECOMMENDED STRATEGY', level=1)

add_instruction_box(doc, "State the recommendation clearly and unambiguously. Describe the implementation mechanism and key conditions for success.")

p = doc.add_paragraph()
run = p.add_run("RECOMMENDATION: ")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)
p.add_run("[State the recommended strategy in bold, clear terms.]")

add_navy_heading(doc, '7.1 Implementation Mechanism', level=2)
doc.add_paragraph("[How will this be executed? What contracts need to be in place? What processes need to change?]")

add_navy_heading(doc, '7.2 Key Conditions', level=2)
doc.add_paragraph("[What must be true for this to succeed? Volume commitments? Executive sponsorship? Timeline constraints?]")

add_navy_heading(doc, '7.3 Dependencies and Risks', level=2)
doc.add_paragraph("[What could derail this? How will those risks be mitigated?]")

doc.add_page_break()

# ============================================
# SECTION 8: BUSINESS CASE
# ============================================
add_navy_heading(doc, '8. BUSINESS CASE', level=1)

add_instruction_box(doc, "Present the full financial case: baseline, value streams, programme costs, scenario analysis, NPV, ROI, and break-even.")

add_navy_heading(doc, '8.1 Baseline and Cost of Inaction', level=2)
doc.add_paragraph("[What is the current total cost of ownership? What are the quantified pain points? What is the cost of doing nothing?]")

add_navy_heading(doc, '8.2 Value Streams', level=2)

# Value streams table
value_table = doc.add_table(rows=10, cols=4)
value_table.style = 'Light Grid Accent 1'
value_headers = ['Value Stream', 'Annual Value (€)', 'Confidence', 'Weighted (€)']
for idx, header in enumerate(value_headers):
    cell = value_table.cell(0, idx)
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)

value_data = [
    ["Cost Reduction", "[Amount]", "[%]", "=B*C"],
    ["Delay Avoidance", "[Amount]", "[%]", "=B*C"],
    ["PPA Revenue", "[Amount]", "[%]", "=B*C"],
    ["Risk Reduction", "[Amount]", "[%]", "=B*C"],
    ["ESG Compliance", "[Amount]", "[%]", "=B*C"],
    ["Working Capital", "[Amount]", "[%]", "=B*C"],
    ["Quality/Performance", "[Amount]", "[%]", "=B*C"],
    ["Innovation Access", "[Amount]", "[%]", "=B*C"],
    ["Strategic Optionality", "[Amount]", "[%]", "=B*C"],
]

for row_idx, row_data in enumerate(value_data, 1):
    for col_idx, value in enumerate(row_data):
        value_table.cell(row_idx, col_idx).text = value

doc.add_paragraph()

add_navy_heading(doc, '8.3 Programme Costs', level=2)

# Costs table
cost_table = doc.add_table(rows=5, cols=4)
cost_table.style = 'Light Grid Accent 1'
cost_headers = ['Cost Category', 'Year 1', 'Annual (Y2-5)', 'Total 5-Year']
for idx, header in enumerate(cost_headers):
    cell = cost_table.cell(0, idx)
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)

cost_data = [
    ["Implementation", "[Amount]", "[Amount]", "[Amount]"],
    ["Annual Running", "[Amount]", "[Amount]", "[Amount]"],
    ["Opportunity Cost", "[Amount]", "[Amount]", "[Amount]"],
    ["TOTAL", "[Sum]", "[Sum]", "[Sum]"],
]

for row_idx, row_data in enumerate(cost_data, 1):
    for col_idx, value in enumerate(row_data):
        cost_table.cell(row_idx, col_idx).text = value

doc.add_paragraph()

add_navy_heading(doc, '8.4 Scenario Analysis', level=2)

# Scenario table
scenario_table = doc.add_table(rows=6, cols=4)
scenario_table.style = 'Light Grid Accent 1'
scenario_headers = ['Metric', 'Bear (70%)', 'Base (100%)', 'Bull (130%)']
for idx, header in enumerate(scenario_headers):
    cell = scenario_table.cell(0, idx)
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)

scenario_data = [
    ["Annual Gross Value", "[70% of base]", "[Base]", "[130% of base]"],
    ["Annual Programme Cost", "[Fixed]", "[Fixed]", "[Fixed]"],
    ["Annual Net Value", "[Gross - Cost]", "[Gross - Cost]", "[Gross - Cost]"],
    ["5-Year NPV (8%)", "[PV calc]", "[PV calc]", "[PV calc]"],
    ["ROI", "[:1]", "[:1]", "[:1]"],
]

for row_idx, row_data in enumerate(scenario_data, 1):
    for col_idx, value in enumerate(row_data):
        scenario_table.cell(row_idx, col_idx).text = value

doc.add_paragraph()

add_navy_heading(doc, '8.5 Financial Summary', level=2)
doc.add_paragraph("[Base case annual net value: €X. 5-year NPV at 8% discount: €Y. ROI: Z:1. Payback: N months. Break-even threshold: M projects per year.]")

doc.add_page_break()

# ============================================
# SECTION 9: IMPLEMENTATION ROADMAP
# ============================================
add_navy_heading(doc, '9. IMPLEMENTATION ROADMAP', level=1)

add_instruction_box(doc, "Provide a phased implementation plan with actions, owners, timelines, and outputs for each phase.")

# Roadmap table
roadmap_table = doc.add_table(rows=9, cols=5)
roadmap_table.style = 'Light Grid Accent 1'
roadmap_headers = ['Phase', 'Actions', 'Owner', 'Timeline', 'Output']
for idx, header in enumerate(roadmap_headers):
    cell = roadmap_table.cell(0, idx)
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)

roadmap_data = [
    ["Phase 0.1", "[Actions]", "[Owner]", "[Weeks]", "[Deliverable]"],
    ["Phase 0.2", "[Actions]", "[Owner]", "[Weeks]", "[Deliverable]"],
    ["Phase 0.3", "[Actions]", "[Owner]", "[Weeks]", "[Deliverable]"],
    ["Phase 1", "[Actions]", "[Owner]", "[Weeks]", "[Deliverable]"],
    ["Phase 2", "[Actions]", "[Owner]", "[Weeks]", "[Deliverable]"],
    ["Phase 3", "[Actions]", "[Owner]", "[Weeks]", "[Deliverable]"],
    ["Phase 4", "[Actions]", "[Owner]", "[Weeks]", "[Deliverable]"],
    ["Phase 5", "[Actions]", "[Owner]", "[Weeks]", "[Deliverable]"],
]

for row_idx, row_data in enumerate(roadmap_data, 1):
    for col_idx, value in enumerate(row_data):
        roadmap_table.cell(row_idx, col_idx).text = value

doc.add_paragraph()

add_navy_heading(doc, '9.1 Key Milestones', level=2)
doc.add_paragraph("[List the critical milestones and decision points. What must be achieved by when?]")

add_navy_heading(doc, '9.2 Resource Requirements', level=2)
doc.add_paragraph("[What internal and external resources are needed? FTE requirements? Budget? External consultants?]")

doc.add_page_break()

# ============================================
# SECTION 10: APPENDICES
# ============================================
add_navy_heading(doc, '10. APPENDICES', level=1)

add_instruction_box(doc, "Include detailed MCDM workings, data sources, and assumptions register.")

add_navy_heading(doc, 'Appendix A: Full MCDM Workings', level=2)
doc.add_paragraph("[AHP pairwise comparison matrix, normalized weights calculation, TOPSIS normalized matrix, weighted normalized matrix, PIS/NIS calculations, distance calculations.]")

add_navy_heading(doc, 'Appendix B: Data Sources', level=2)
doc.add_paragraph("[List all data sources used: company reports, industry databases, news sources, indices, expert interviews.]")

add_navy_heading(doc, 'Appendix C: Assumptions Register', level=2)

# Assumptions table
assumptions_table = doc.add_table(rows=6, cols=4)
assumptions_table.style = 'Light Grid Accent 1'
assumptions_headers = ['ID', 'Assumption', 'Rationale', 'Impact if Wrong']
for idx, header in enumerate(assumptions_headers):
    cell = assumptions_table.cell(0, idx)
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x21, 0x47)

assumptions_data = [
    ["A1", "[Assumption text]", "[Why assumed]", "[Impact]"],
    ["A2", "[Assumption text]", "[Why assumed]", "[Impact]"],
    ["A3", "[Assumption text]", "[Why assumed]", "[Impact]"],
    ["A4", "[Assumption text]", "[Why assumed]", "[Impact]"],
    ["A5", "[Assumption text]", "[Why assumed]", "[Impact]"],
]

for row_idx, row_data in enumerate(assumptions_data, 1):
    for col_idx, value in enumerate(row_data):
        assumptions_table.cell(row_idx, col_idx).text = value

# Save document
output_path = '/Users/jonathonmilne/.openclaw/workspace/category-strategy-process/CATEGORY_STRATEGY_TEMPLATE.docx'
doc.save(output_path)
print(f"Word template saved: {output_path}")
